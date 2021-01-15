# -*- coding: utf-8 -*-

'''A group of ``TextBlock``, ``ImageBlock`` or ``TableBlock``.
'''

from docx.shared import Pt
from ..common import constants
from ..common.Collection import Collection
from ..common.share import BlockType, rgb_value
from ..common.Block import Block
from ..common.docx import reset_paragraph_format
from ..text.TextBlock import TextBlock
from ..text.Line import Line
from ..text.Lines import Lines
from ..image.ImageBlock import ImageBlock
from ..table.TableBlock import TableBlock
from ..table.Cell import Cell


class Blocks(Collection):
    '''Block collections.'''
    def __init__(self, instances:list=None, parent=None):
        ''' A collection of TextBlock and TableBlock instances. 
            ImageBlock is converted to ImageSpan contained in TextBlock.
        '''
        self._parent = parent # type: Layout
        self._instances = []  # type: list[TextBlock or TableBlock]

        # NOTE: no changes on floating image blocks once identified, so store them here
        self.floating_image_blocks = []
    
        # Convert all original image blocks to text blocks, i.e. ImageSpan,
        # So we can focus on single TextBlock later on; TableBlock is also combination of TextBlocks.
        # NOTE: for a converted image block
        # - isinstance(image_block, TextBlock)==True
        # - image_block.is_text_block()==False
        for block in (instances or []):
            if isinstance(block, ImageBlock):
                text_block = block.to_text_block()
                self.append(text_block)
            else:
                self.append(block)


    def _update_bbox(self, block:Block):
        ''' Override. The parent is generally ``Layout`` or ``Cell``, which is not necessary to 
            update its bbox. So, do nothing but required here.
        '''
        pass


    @property
    def lattice_table_blocks(self):
        '''Get lattice table blocks contained in this Collection.'''
        return list(filter(
            lambda block: block.is_lattice_table_block(), self._instances))


    @property
    def stream_table_blocks(self):
        '''Get stream table blocks contained in this Collection.'''
        return list(filter(
            lambda block: block.is_stream_table_block(), self._instances))


    @property
    def table_blocks(self):
        '''Get table blocks contained in this Collection.'''
        return list(filter(
            lambda block: block.is_table_block(), self._instances))


    @property
    def inline_image_blocks(self):
        '''Get inline image blocks contained in this Collection.'''
        return list(filter(
            lambda block: block.is_inline_image_block(), self._instances))


    @property
    def text_blocks(self):
        '''Get text/image blocks contained in this Collection.'''
        return list(filter(
            lambda block: block.is_text_image_block(), self._instances))

    
    def store(self):
        '''Store attributes in json format.'''
        res = super().store()
        res.extend([ instance.store() for instance in self.floating_image_blocks ])
        return res


    def restore(self, raws:list):
        '''Clean current instances and restore them from source dict.

        Args:
            raws (list): A list of raw dicts representing text/image/table blocks.

        Returns:
            Blocks: self
        '''
        self.reset()  # clean current instances
        for raw_block in raws:
            block_type = raw_block.get('type', -1) # type: int
            
            # inline image block -> text block
            if block_type==BlockType.IMAGE.value:
                block = ImageBlock(raw_block).to_text_block()
            
            # text block
            elif block_type == BlockType.TEXT.value:
                block = TextBlock(raw_block)
            
            # floating image block
            elif block_type == BlockType.FLOAT_IMAGE.value:
                block = ImageBlock(raw_block)
                block.set_float_image_block()

            # table block
            elif block_type in (BlockType.LATTICE_TABLE.value, BlockType.STREAM_TABLE.value):
                block = TableBlock(raw_block)
            
            else:
                block = None            
            
            # add to list
            self.append(block)
        
        return self


    def clean_up(self, float_image_ignorable_gap:float, line_overlap_threshold:float, line_merging_threshold:float):
        """Clean up blocks in page level.

        * remove blocks out of page
        * remove transformed text: text direction is not (1, 0) or (0, -1)
        * remove empty blocks

        Args:
            float_image_ignorable_gap (float): Regarded as float image if the intersection exceeds this value.
            line_overlap_threshold (float): Delete line if the intersection to other lines exceeds this value.
            line_merging_threshold (float): Combine two lines if the x-distance is lower than this value.

        .. note::
            This method works ONLY for layout initialized from raw dict extracted by ``page.getText()``.
            Under this circumstance, it only exists text blocks since all raw image blocks are converted to 
            text blocks.
        """
        if not self._instances: return
        
        page_bbox = self.parent.bbox
        f = lambda block:   block.bbox.intersects(page_bbox) and \
                            block.text.strip() and (
                            block.is_horizontal_text or block.is_vertical_text)
        self.reset(filter(f, self._instances))

        # merge blocks horizontally, e.g. remove overlap blocks, since no floating elements are supported
        # NOTE: It's to merge blocks in physically horizontal direction, i.e. without considering text direction.
        self.strip() \
            .sort_in_reading_order() \
            .identify_floating_images(float_image_ignorable_gap) \
            .join_horizontally(text_direction=False, 
                            line_overlap_threshold=line_overlap_threshold,
                            line_merging_threshold=line_merging_threshold)
   
    def strip(self):
        '''Remove redundant blanks exist in text block lines. These redundant blanks may affect bbox of text block.
        '''
        for block in self._instances: block.strip()
        return self


    def identify_floating_images(self, float_image_ignorable_gap:float):
        """Identify floating image lines and convert to ImageBlock.

        Args:
            float_image_ignorable_gap (float): Regarded as float image if the intersection exceeds this value.
        """
        # group lines by connectivity
        lines = Lines()
        for block in self._instances:
            lines.extend(block.lines)
        groups = lines.group_by_connectivity(dx=-float_image_ignorable_gap, dy=-float_image_ignorable_gap)
        
        # identify floating objects
        for group in filter(lambda group: len(group)>1, groups):
            for line in group:
                block = line.parent
                # consider image block only (converted to text block)
                if not block.is_inline_image_block(): continue

                float_image = ImageBlock().from_text_block(block)
                float_image.set_float_image_block()
                self.floating_image_blocks.append(float_image)

                # remove the original image block from flow layout
                block.update_bbox((0,0,0,0))

        return self


    def assign_to_tables(self, tables:list):
        """Add Text/Image/table block lines to associated cells of given tables.

        Args:
            tables (list): A list of TableBlock instances.
        """        
        if not tables: return

        # assign blocks to table region        
        blocks_in_tables = [[] for _ in tables] # type: list[list[Block]]
        blocks = []   # type: list[Block]
        for block in self._instances:
            self._assign_block_to_tables(block, tables, blocks_in_tables, blocks)

        # assign blocks to associated cells
        for table, blocks_in_table in zip(tables, blocks_in_tables):
            # no contents for this table
            if not blocks_in_table: continue
            table.assign_blocks(blocks_in_table)

        # sort in natural reading order and update layout blocks
        blocks.extend(tables)
        self.reset(blocks).sort_in_reading_order()        


    def collect_stream_lines(self, potential_shadings:list, 
            float_layout_tolerance:float, 
            line_separate_threshold:float):
        '''Collect elements in Line level, which may contained in a stream table region.

        * Lines in text block
        * The entire bbox of table block
        
        Table may exist on the following conditions:

        * (a) lines in potential shading -> determined by shapes
        * (b) lines in blocks are not connected sequently -> determined by current block only
        * (c) multi-blocks are in a same row (horizontally aligned) -> determined by two adjacent blocks
        
        Args:
            potential_shadings (list): a group of shapes representing potential cell shading
        
        Returns:
            list: A list of Lines. Each group of Lines represents a potential table.
        
        .. note::
            ``PyMuPDF`` may group multi-lines in a row as a text block while each line belongs to different
            cell. So, it's required to deep into line level.
        '''
        # get sub-lines from block
        def sub_lines(block):
            return block.lines if block.is_text_image_block() else [Line().update_bbox(block.bbox)]
        
        # exclude potential shading in white bg-color
        shadings_exclude_white = list(filter(
            lambda shape: shape.color != rgb_value((1,1,1)), potential_shadings
        ))
        
        # check block by block
        res = [] # type: list[Lines]
        j = 0
        table_lines = Lines() # potential text lines in a table
        new_line = True
        num_blocks, num_shadings = len(self._instances), len(shadings_exclude_white)
        for i in range(num_blocks):

            block = self._instances[i]
            next_block = self._instances[i+1] if i<num_blocks-1 else Block()

            table_end = False

            # (a) block in potential shading?
            if j < num_shadings:
                shading = shadings_exclude_white[j]
                if not shading.is_determined and shading.contains(block, threshold=constants.FACTOR_MOST):
                    table_lines.extend(sub_lines(block))
                    new_line = False
                
                # move to next shading
                elif next_block.bbox.y0 > shading.bbox.y1:
                    j += 1
            
            # (b) add lines when current block is not flow layout
            if new_line and not block.is_flow_layout(float_layout_tolerance, line_separate_threshold): 
                table_lines.extend(sub_lines(block))  # deep into line level
                
                # update line status
                new_line = False            

            # (c) multi-blocks are in a same row: check layout with next block?
            # yes, add both current and next blocks
            if block.horizontally_align_with(next_block, factor=0):
                # if it's start of new table row: add the first block
                if new_line: table_lines.extend(sub_lines(block))
                
                # add next block
                table_lines.extend(sub_lines(next_block))

                # update line status
                new_line = False

            # no, consider to start a new row
            else:
                # table end if it's a text line, i.e. no more than one block in a same line
                if new_line: table_end = True

                # update line status            
                new_line = True

            # NOTE: close table detecting manually if last block
            if i==num_blocks-1: table_end = True

            # end of current table
            if table_lines and table_end: 
                res.append(table_lines)                
                table_lines = Lines() # reset table_blocks

        return res


    def parse_spacing(self, *args):
        '''Calculate external and internal vertical space for text blocks.
        
        * Paragraph spacing is determined by the vertical distance to previous block.
          For the first block, the reference position is top margin.
        * Line spacing is defined as the average line height in current block.

        It's easy to set before-space or after-space for a paragraph with ``python-docx``,
        so, if current block is a paragraph, set before-space for it; if current block is 
        not a paragraph, e.g. a table, set after-space for previous block (generally, 
        previous block should be a paragraph).
        '''
        if not self._instances: return

        # bbox of blocks
        # - page level, e.g. blocks in top layout
        # - table level, e.g. blocks in table cell
        bbox = self.parent.bbox

        # check text direction for vertical space calculation:
        # - normal reading direction (from left to right)    -> the reference boundary is top border, i.e. bbox[1].
        # - vertical text direction, e.g. from bottom to top -> left border bbox[0] is the reference
        idx = 1 if self.is_horizontal_text else 0

        ref_block = self._instances[0]
        ref_pos = bbox[idx]

        for block in self._instances:

            #---------------------------------------------------------
            # alignment mode and left spacing:
            # - horizontal block -> take left boundary as reference
            # - vertical block   -> take bottom boundary as reference
            #---------------------------------------------------------
            block.parse_horizontal_spacing(bbox, *args)            

            #---------------------------------------------------------
            # vertical space calculation
            #---------------------------------------------------------

            # NOTE: the table bbox is counted on center-line of outer borders, so a half of top border
            # size should be excluded from the calculated vertical spacing
            if block.is_table_block():
                dw = block[0][0].border_width[0] / 2.0 # use top border of the first cell
            
            else:
                dw = 0.0

            start_pos = block.bbox[idx] - dw
            para_space = start_pos-ref_pos

            # modify vertical space in case the block is out of bottom boundary
            dy = max(block.bbox[idx+2]-bbox[idx+2], 0.0)
            para_space -= dy
            para_space = max(para_space, 0.0) # ignore negative value

            # ref to current (paragraph): set before-space for paragraph
            if block.is_text_block():

                # spacing before this paragraph
                block.before_space = start_pos-ref_pos # keep negative value temperally

                # calculate average line spacing in paragraph
                # NOTE: adjust before space if negative value
                block.parse_line_spacing()

            # if ref to current (image): set before-space for paragraph
            elif block.is_inline_image_block():
                block.before_space = para_space

            # ref (paragraph/image) to current: set after-space for ref paragraph        
            elif ref_block.is_text_block() or ref_block.is_inline_image_block():
                ref_block.after_space = para_space

            # situation with very low probability, e.g. ref (table) to current (table)
            # we can't set before space for table in docx, but the tricky way is to
            # create an empty paragraph and set paragraph line spacing and before space
            else:
                # let para_space>=1 Pt to accommodate the dummy paragraph if not the first block
                block.before_space = max(para_space, int(block!=self._instances[0])*constants.MINOR_DIST)

            # update reference block        
            ref_block = block
            ref_pos = ref_block.bbox[idx+2] + dw # assume same bottom border with top one

        # NOTE: when a table is at the end of a page, a dummy paragraph with a small line spacing 
        # is added after this table, to avoid unexpected page break. Accordingly, this extra spacing 
        # must be remove in other place, especially the page space is run out.
        # Here, reduce the last row of table.
        block = self._instances[-1]
        if block.is_table_block(): block[-1].height -= constants.MINOR_DIST


    def join_horizontally(self, 
                text_direction:bool, 
                line_overlap_threshold:float, 
                line_merging_threshold:float):
        '''Join lines in horizontally aligned blocks into new TextBlock.

        This function converts potential float layout into flow layout, e.g. 
        reposition inline images, so that make rebuilding such layout in docx possible.
        
        Args:
            text_direction (bool): Whether consider text direction. 
                                   Detect text direction based on line direction if True, 
                                   otherwise use default direction, i.e. from left to right.
        '''
        # get horizontally aligned blocks group by group
        fun = lambda a,b: a.horizontally_align_with(b, factor=0.0, text_direction=text_direction)
        groups = self.group(fun)
        
        # merge text blocks in each group
        blocks = []
        for blocks_collection in groups:
            # combine all lines into a TextBlock
            final_block = TextBlock()
            for block in blocks_collection:
                if not block.is_text_image_block(): continue
                final_block.add(block.lines) # keep empty line, may help to identify table layout

            # merge lines/spans contained in this text block
            # NOTE:            
            # Lines in text block must have same property, e.g. height, vertical distance, 
            # because average line height is used when create docx. However, the contained lines 
            # may be not reasonable after this step. So, this is just a pre-processing step focusing 
            # on processing lines in horizontal direction, e.g. merging inline image to its text line.
            # Further steps, e.g. split back to original blocks, must be applied before further parsing.
            final_block.lines.join(line_overlap_threshold, line_merging_threshold)

            blocks.append(final_block)
        
        # add table blocks
        blocks.extend(self.table_blocks)
        self.reset(blocks)

        return self


    def split_back(self, *args):
        '''Split the joined lines back to original text block if possible.

        With preceding joining step, current text block may contain lines coming from various original blocks.
        Considering that different text block may have different line properties, e.g. height, spacing, 
        this function is to split them back to original text block. 

        .. note::
            Don't split block if the splitting breaks flow layout, e.g. two blocks (i.e. two paragraphs in docx) 
            in same row.
        '''
        blocks = [] # type: list[TextBlock]
        lines = Lines()
        # collect lines for further step, or table block directly
        for block in self._instances:
            if block.is_text_image_block() and block.is_flow_layout(*args):
                lines.extend(block.lines)
            else:
                blocks.append(block)
        
        # regroup lines
        for group_lines in lines.split_back():
            text_block = TextBlock()
            text_block.lines.reset(group_lines)
            blocks.append(text_block)
       
        self.reset(blocks).sort_in_reading_order()


    def parse_text_format(self, rects):
        '''Parse text format with style represented by stroke/fill shapes.
        
        Args:
            rects (Shapes): Potential styles applied on blocks.
        '''
        # parse text block style one by one
        for block in filter(lambda e: e.is_text_block(), self._instances): 
            block.parse_text_format(rects)


    def make_docx(self, doc):
        '''Create page based on parsed block structure. 
        
        Args:
            doc (Document, _Cell): The container to make docx content.
        '''
        def make_table(table_block, pre_table):
            # create dummy paragraph if table before space is set
            # - a minimum line height of paragraph is 0.7pt, so ignore before space if less than this value
            # - but tow adjacent tables will be combined automatically, so adding a minimum dummy paragraph is required
            if table_block.before_space>=constants.MIN_LINE_SPACING or pre_table:
                h = int(10*table_block.before_space)/10.0 # round(x,1), but to lower bound
                p = doc.add_paragraph()
                reset_paragraph_format(p, line_spacing=Pt(h))

            # new table            
            table = doc.add_table(rows=table_block.num_rows, cols=table_block.num_cols)
            table.autofit = False
            table.allow_autofit  = False
            table_block.make_docx(table)

        pre_table = False
        for block in self._instances:
            # make paragraphs
            if block.is_text_image_block():
                # new paragraph
                p = doc.add_paragraph()
                block.make_docx(p)

                pre_table = False # mark block type
            
            # make table
            elif block.is_table_block():
                make_table(block, pre_table)
                pre_table = True # mark block type

        # below table processing is necessary for page level only
        if isinstance(self.parent, Cell): return
        
        # NOTE: If a table is at the end of a page, a new paragraph will be automatically 
        # added by the rending engine, e.g. MS Word, which resulting in an unexpected
        # page break. The solution is to never put a table at the end of a page, so add
        # an empty paragraph and reset its format, particularly line spacing, when a table
        # is created.
        for block in self._instances[::-1]:
            # ignore float image block
            if block.is_float_image_block(): continue

            # nothing to do if not end with table block
            if not block.is_table_block(): break

            # otherwise, add a small paragraph
            p = doc.add_paragraph()
            reset_paragraph_format(p, Pt(constants.MIN_LINE_SPACING)) # a small line height

        # Finally, add floating image to last paragraph
        p = doc.add_paragraph() if not doc.paragraphs else doc.paragraphs[-1]
        for image_block in self.floating_image_blocks:
            image_block.make_docx(p)

  
    def plot(self, page):
        '''Plot blocks in PDF page for debug purpose.'''
        for block in self._instances: block.plot(page)                


    @staticmethod
    def _assign_block_to_tables(block:Block, tables:list, blocks_in_tables:list, blocks:list):
        '''Collect blocks contained in table region ``blocks_in_tables`` and rest text blocks in ``blocks``.'''
        # lines in block for further check if necessary
        lines = block.lines if block.is_text_image_block() else Lines()

        # collect blocks contained in table region
        # NOTE: there is a probability that only a part of a text block is contained in table region, 
        # while the rest is in normal block region.
        for table, blocks_in_table in zip(tables, blocks_in_tables):

            # fully contained in one table
            if table.bbox.contains(block.bbox):
                blocks_in_table.append(block)
                break
            
            # not possible in current table, then check next table
            elif not table.bbox.intersects(block.bbox): continue
            
            # deep into line level for text block
            elif block.is_text_image_block():
                table_lines, not_table_lines = lines.split_with_intersection(table.bbox, constants.FACTOR_MOST)

                # add lines to table
                text_block = TextBlock()
                text_block.add(table_lines)
                blocks_in_table.append(text_block)

                # lines not in table for further check
                if not_table_lines:
                    lines = not_table_lines
                else:
                    break # no more lines
        
        # Now, this block (or part of it) belongs to previous layout
        else:
            if block.is_table_block():
                blocks.append(block)
            else:
                text_block = TextBlock()
                text_block.add(lines)
                blocks.append(text_block)


