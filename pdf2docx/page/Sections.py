# -*- coding: utf-8 -*-

'''Collection of :py:class:`~pdf2docx.page.Section` instances.
'''

from docx.enum.section import WD_SECTION
from docx.shared import Pt
from ..common.Collection import Collection, BaseCollection
from ..common.Block import Block
from ..common.docx import reset_paragraph_format
from ..layout.Blocks import Blocks
from ..shape.Shapes import Shapes
from ..shape.Shape import Shape
from .Section import Section
from .Column import Column


class Sections(BaseCollection):

    def restore(self, raws:list):
        """Restore sections from source dicts."""        
        self.reset()
        for raw in raws:
            section = Section().restore(raw)
            self.append(section)
        return self
    

    def parse(self, settings:dict):
        '''Parse sections layout.'''
        # detect section based on source blocks and shapes
        blocks, shapes = self.parent.blocks, self.parent.shapes
        self._parse(blocks, shapes)

        # clear source blocks and shapes once created sections
        blocks.reset()
        shapes.reset()

        # parse layout under section level
        for section in self: section.parse(settings)        
        return self


    def make_docx(self, doc):
        '''Create sections in docx.'''
        # -----------------
        # new page
        # -----------------
        if doc.paragraphs:
            doc.add_section(WD_SECTION.NEW_PAGE)

        # -----------------
        # first section
        # -----------------
        # vertical position
        p = doc.add_paragraph()
        line_height = min(self[0].before_space, 11)
        pf = reset_paragraph_format(p, line_spacing=Pt(line_height))
        pf.space_after = Pt(self[0].before_space-line_height)
        if self[0].cols==2: 
            doc.add_section(WD_SECTION.CONTINUOUS)
        
        # create first section
        self[0].make_docx(doc)        

        # -----------------
        # more sections
        # -----------------
        for section in self[1:]:
            # create new section symbol
            doc.add_section(WD_SECTION.CONTINUOUS)

            # set after space of last paragraph to define the vertical
            # position of current section
            p = doc.paragraphs[-2] # -1 is the section break
            pf = p.paragraph_format
            pf.space_after = Pt(section.before_space)
            
            # section content
            section.make_docx(doc)


    def plot(self, page):
        '''Plot all section blocks for debug purpose.'''
        for section in self: 
            for column in section:
                column.blocks.plot(page)
    

    def _parse(self, blocks:Blocks, shapes:Shapes):
        '''Detect and create page sections based on initial layout extracted from ``PyMuPDF``.

        Args:
            blocks (Blocks): source blocks after cleaning up.
            shapes (SHapes): source shapes after cleaning up.

        .. note::
            Consider two-columns Section only.
        '''
        # bbox
        X0, Y0, X1, _ = self.parent.working_bbox        
    
        # collect all blocks and shapes
        elements = Collection()
        elements.extend(blocks)
        elements.extend(shapes)
        
        # check section row by row
        pre_section = Collection()
        pre_num_col = 1
        y_ref = Y0 # to calculate v-distance between sections
        for row in elements.group_by_rows():
            # check column col by col
            cols = row.group_by_columns()
            current_num_col = len(cols)

            # consider 2-cols only
            if current_num_col>2: current_num_col = 1 

            # further check 2-cols -> the height
            x0, y0, x1, y1 = pre_section.bbox
            if pre_num_col==2 and current_num_col==1 and y1-y0<20:
                pre_num_col = 1

            # TODO:
            # 1. pre_num_col==2 and current_num_col==1, but current row in the left side
            # 2. pre_num_col==2 and current_num_col==2, but not aligned

            # finalize pre-section if different to current section
            if current_num_col!=pre_num_col:
                # create pre-section
                section = self._create_section(pre_num_col, pre_section, (X0, X1), y_ref)
                if section:
                    self.append(section)
                    y_ref = section[-1].bbox[3]

                # start new section                
                pre_section = Collection(row)
                pre_num_col = current_num_col

            # otherwise, append to pre-section
            else:
                pre_section.extend(row)

        # the final section
        section = self._create_section(current_num_col, pre_section, (X0, X1), y_ref)
        self.append(section)


    @staticmethod
    def _create_column(bbox, elements:Collection):
        '''Create column based on bbox and candidate elements: blocks and shapes.'''
        if not bbox: return None
        column = Column().update_bbox(bbox)
        blocks = [e for e in elements if isinstance(e, Block)]
        shapes = [e for e in elements if isinstance(e, Shape)]
        column.assign_blocks(blocks)
        column.assign_shapes(shapes)
        return column


    @staticmethod
    def _create_section(num_col:int, elements:Collection, h_range:tuple, y_ref:float):
        '''Create section based on column count, candidate elements and horizontal boundary.'''
        if not elements: return
        X0, X1 = h_range

        if num_col==1:
            x0, y0, x1, y1 = elements.bbox
            column = Sections._create_column((X0, y0, X1, y1), elements)
            section = Section(space=0, columns=[column])
            before_space = y0 - y_ref
        else:
            cols = elements.group_by_columns()
            u0, v0, u1, v1 = cols[0].bbox
            m0, n0, m1, n1 = cols[1].bbox
            u = (u1+m0)/2.0
            column_1 = Sections._create_column((X0, v0, u, v1), elements)
            column_2 = Sections._create_column((u, n0, X1, n1), elements)
            section = Section(space=0, columns=[column_1, column_2])
            before_space = v0 - y_ref

        section.before_space = round(before_space, 1)
        return section
                