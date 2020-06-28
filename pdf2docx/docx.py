'''
create *.docx file based on PDF layout data with python package python-docx.
@created: 2019-06-28
@author: train8808@gmail.com
'''


from io import BytesIO

from docx.shared import Pt
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ROW_HEIGHT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from . import utils
from .pdf_block import (is_text_block, is_image_block, is_table_block, is_explicit_table_block)



def make_page(doc, layout):
    ''' create page based on layout data. 

        To avoid incorrect page break from original document, a new page section
        is created for each page.

        Support general document style only:
          - writing mode: from left to right, top to bottom
          - text direction: horizontal

        The vertical postion of paragraph/table is defined by space_before or 
        space_after property of a paragraph.
    '''
    # new page section
    # a default section is created when initialize the document,
    # so we do not have to add section for the first time.
    if not doc.paragraphs:
        section = doc.sections[0]
    else:
        section = doc.add_section(WD_SECTION.NEW_PAGE)

    width, height = layout['width'], layout['height']
    section.page_width  = Pt(width)
    section.page_height = Pt(height)

    # set page margin
    left,right,top,bottom = layout['margin']    
    section.left_margin = Pt(left)
    section.right_margin = Pt(right)
    section.top_margin = Pt(top)
    section.bottom_margin = Pt(bottom)

    # add paragraph or table according to parsed block
    for block in layout['blocks']:
        # make paragraphs
        if is_text_block(block) or is_image_block(block):
            # horizontal paragraph
            if is_image_block(block) or block['lines'][0]['wmode'] == 0:
                # new paragraph
                p = doc.add_paragraph()
                left, right, *_ = layout['margin']
                make_paragraph(p, block, left, width-right)
            
            # vertical paragraph: TODO
            else:
                pass
        
        # make table
        elif is_table_block(block):
            # new table            
            table = doc.add_table(rows=len(block['cells']), cols=len(block['cells'][0]))
            table.autofit = False
            table.allow_autofit  = False
            make_table(table, block, width, layout['margin'])
            
            # NOTE: If this table is at the end of a page, a new paragraph is automatically 
            # added by the rending engine, e.g. MS Word, which resulting in an unexpected
            # page break. The solution is to never put a table at the end of a page, so add
            # an empty paragraph and reset its format, particularly line spacing, when a table
            # is created.
            p = doc.add_paragraph()
            _reset_paragraph_format(p, Pt(1.0))


def make_paragraph(p, block, X0, X1):
    '''create paragraph for a text block.
       join line sets with TAB and set position according to bbox.

       Generally, a pdf block is a docx paragraph, with block|line as line in paragraph.
       But without the context, it's not able to recognize a block line as word wrap, or a 
       separate line instead. A rough rule used here:
        - block line will be treated as separate line (append `\n`) by default, except
        - (1) this line and next line are actually in the same line (y-position)
        - (2) if the rest space of this line can't accommodate even one span of next line, 
              it's supposed to be normal word wrap.

        ---
        Args:
            X0, X1: the paragraph is restricted in a horizontal range within (X0, X1)
    '''
    # indent and space setting
    before_spacing = max(round(block.get('before_space', 0.0), 1), 0.0)
    after_spacing = max(round(block.get('after_space', 0.0), 1), 0.0)
    pf = _reset_paragraph_format(p)
    pf.space_before = Pt(before_spacing)
    pf.space_after = Pt(after_spacing)    

    # add image
    if is_image_block(block):
        # left indent implemented with tab
        pos = block['bbox'][0]-X0
        if abs(pos) > utils.DM:
            pf.tab_stops.add_tab_stop(Pt(pos))
            p.add_run().add_tab()
        # create image with bytes data stored in block.
        span = p.add_run()
        span.add_picture(BytesIO(block['image']), width=Pt(block['bbox'][2]-block['bbox'][0]))

    # add text (inline image may exist)
    else:
        # set line spacing for text paragraph
        pf.line_spacing = Pt(round(block['line_space'],1))

        for i, line in enumerate(block['lines']):

            # left indent implemented with tab
            pos = line['bbox'][0]-X0
            if abs(pos) > utils.DM:
                pf.tab_stops.add_tab_stop(Pt(pos))
                p.add_run().add_tab()

            # add line
            for span in line['spans']:
                # add content
                _add_span(span, p)

                # exactly line spacing will destroy image display, so set single line spacing instead
                if 'image' in span:
                    pf.line_spacing = 1.05

            # break line or word wrap?
            # new line by default
            line_break = True

            # no more lines after last line
            if line==block['lines'][-1]: 
                line_break = False
            
            # do not break line if they're indeed in same line
            elif utils.in_same_row(block['lines'][i+1]['bbox'], line['bbox']):
                line_break = False
            
            # now, we have two lines, check whether word wrap or line break
            else:                
                # word wrap if rest space of this line can't accommodate even one word of next line span
                # bbox of first span in next line
                next_span = block['lines'][i+1]['spans'][0]
                required_space = next_span['bbox'][2] - next_span['bbox'][0] # width of whole span

                # calculate first word length approximately for text span
                if 'image' not in next_span:                    
                    words = next_span['text'].strip()
                    if words:
                        word = words.split(' ')[0]
                        required_space *= len(word) / len(words)

                free_space = X1-line['bbox'][2]
                if required_space >= free_space:
                    line_break = False
            
            if line_break:
                p.add_run('\n')

    return p
    

def make_table(table, block, page_width, page_margin):
    '''create table for a text block
       count of columns are checked, combine rows if next block is also in table format
    '''
    # set indent    
    pos = block['bbox'][0]-page_margin[0]
    _indent_table(table, pos)

    # cell format and contents
    block_cells = block['cells']
    border_style = is_explicit_table_block(block)
    for i in range(len(table.rows)):
        for j in range(len(table.columns)):           

            # ignore merged cells
            block_cell = block_cells[i][j]
            if not block_cell: continue
            
            # set cell style
            # no borders for implicit table
            _set_cell_style(table, (i,j), block_cell, border_style)

            # clear cell margin
            # NOTE: the start position of a table is based on text in cell, rather than left border of table. 
            # They're almost aligned if left-margin of cell is zero.
            cell = table.cell(i, j)
            _set_cell_margins(cell, start=0, end=0)

            # insert text            
            first = True
            x0, _, x1, _ = block_cell['bbox']
            for block in block_cell['blocks']:
                if first:
                    p = cell.paragraphs[0]
                    first = False
                else:
                    p = cell.add_paragraph()
                make_paragraph(p, block, x0, x1)


def _reset_paragraph_format(p, line_spacing=1.05):
    ''' reset paragraph format, especially line spacing.
        Two kinds of line spacing, corresponding to the setting in MS Office Word:
        - line_spacing=1.05: single or multiple
        - line_spacing=Pt(1): exactly
    '''
    pf = p.paragraph_format
    pf.line_spacing = line_spacing # single by default
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.widow_control = True
    return pf


def _add_span(span, paragraph):
    '''add text span to a paragraph.       
    '''
    # inline image span
    if 'image' in span:
        image_span = paragraph.add_run()
        image_span.add_picture(BytesIO(span['image']), width=Pt(span['bbox'][2]-span['bbox'][0]))

    # text span
    else:
        text_span = paragraph.add_run(span['text'])

        # style setting
        # https://python-docx.readthedocs.io/en/latest/api/text.html#docx.text.run.Font

        # basic font style
        # line['flags'] is an integer, encoding bool of font properties:
        # bit 0: superscripted (2^0)
        # bit 1: italic (2^1)
        # bit 2: serifed (2^2)
        # bit 3: monospaced (2^3)
        # bit 4: bold (2^4)            
        text_span.italic = bool(span['flags'] & 2**1)
        text_span.bold = bool(span['flags'] & 2**4)
        text_span.font.name = utils.parse_font_name(span['font'])
        text_span.font.size = Pt(round(span['size']*2)/2.0) # only x.0 and x.5 is accepted in docx
        text_span.font.color.rgb = RGBColor(*utils.RGB_component(span['color']))

        # font style parsed from PDF rectangles: 
        # e.g. highlight, underline, strike-through-line
        for style in span.get('style', []):
            t = style['type']
            if t==0:
                text_span.font.highlight_color = utils.to_Highlight_color(style['color'])
            elif t==1:
                text_span.font.underline = True
            elif t==2:
                text_span.font.strike = True


def _set_cell_style(table, indexes, block_cell, border_style=True):
    ''' Set python-docx cell style, e.g. border, shading, width, row height, 
        based on cell block parsed from PDF.

        Args:
            table: python-docx table object
            indexes: (i, j) index of current cell in table
            block_cell: a list of list, cells information parsed from PDF
    '''
    i, j = indexes
    cell = table.cell(i, j)
    n_row, n_col = block_cell['merged-cells']

    # ---------------------
    # border style
    # ---------------------
    # NOTE: border width is specified in eighths of a point, with a minimum value of 
    # two (1/4 of a point) and a maximum value of 96 (twelve points)
    if border_style:
        keys = ('top', 'end', 'bottom', 'start')
        kwargs = {}
        for k, w, c in zip(keys, block_cell['border-width'], block_cell['border-color']):
            hex_c = f'#{hex(c)[2:].zfill(6)}'
            kwargs[k] = {
                'sz': 8*w, 'val': 'single', 'color': hex_c.upper()
            }
        # merged cells are assumed to have same borders with the main cell        
        for m in range(i, i+n_row):
            for n in range(j, j+n_col):
                _set_cell_border(table.cell(m, n), **kwargs)

    # ---------------------
    # merge cells
    # ---------------------        
    if n_row*n_col!=1:
        _cell = table.cell(i+n_row-1, j+n_col-1)
        cell.merge(_cell)

    # ---------------------
    # cell width/height
    # ---------------------
    x0, y0, x1, y1 = block_cell['bbox']
    
    # set cell height by setting row height
    # NOTE: consider separate rows (without cell merging) only since merged rows are determined accordingly.
    if n_row==1:
        row = table.rows[i]
        # to control the layout precisely, set `exact` value, rather than `at least` value
        # the associated steps in MS word: Table Properties -> Row -> Row height -> exactly
        row.height_rule = WD_ROW_HEIGHT.EXACTLY        
        # NOTE: cell height is counted from center-line of top border to center line of bottom border,
        # i.e. the height of cell bbox
        row.height = Pt(y1-y0) # Note cell does not have height property.    
    
    # set cell width
    # experience: width of merged cells may change if not setting width for merged cells
    cell.width = Pt(x1-x0)

    # ---------------------
    # cell bg-color
    # ---------------------
    if block_cell['bg-color']!=None:
        _set_cell_shading(cell, block_cell['bg-color'])


def _indent_table(table, indent):
    '''indent table

       args:
         - indent: indent value, the basic unit is 1/20 pt
    '''
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(20*indent)) # basic unit 1/20 pt for openxml 
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)


def _set_cell_margins(cell, **kwargs):
    ''' Set cell margins. Provided values are in twentieths of a point (1/1440 of an inch).

        ---
        Args:
            - cell:  actual cell instance you want to modify
            - kwargs: a dict with keys: top, bottom, start, end
        usage:
            _set_cell_margins(cell, top=50, start=50, bottom=50, end=50)
        read more here: 
            - https://blog.csdn.net/weixin_44312186/article/details/104944773
            - http://officeopenxml.com/WPtableCellMargins.php
    '''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
 
    for m in ['top', 'start', 'bottom', 'end']:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
 
    tcPr.append(tcMar)


def _set_cell_shading(cell, RGB_value):
    '''set cell background-color'''
    c = hex(RGB_value)[2:].zfill(6)
    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), c)))


def _set_cell_border(cell, **kwargs):
    """
    Set cell`s border.
    
    Reference:
     - https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx
     - https://blog.csdn.net/weixin_44312186/article/details/104944110

    Usage:

    _set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

