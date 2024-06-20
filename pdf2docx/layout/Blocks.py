'''
A group of text elements, distinguished to ``Shape`` elements. For instance, ``TextBlock``, 
``ImageBlock`` or ``TableBlock`` after parsing, while ``Line`` instances at the beginning, 
and a combination of ``Line`` and ``TableBlock`` during parsing process.
'''

import logging
from docx.shared import Pt
from ..common import constants
from ..common.Collection import ElementCollection
from ..common.share import (BlockType, lower_round, rgb_value)
from ..common.Block import Block
from ..common.docx import (reset_paragraph_format, delete_paragraph)
from ..text.TextBlock import TextBlock
from ..text.TextSpan import TextSpan
from ..text.Line import Line
from ..text.Lines import Lines
from ..table.Cell import Cell
from ..image.ImageBlock import ImageBlock
from ..table.TableBlock import TableBlock


class Blocks(ElementCollection):
    '''Block collections.'''
    def __init__(self, instances:list=None, parent=None):
        ''' A collection of text based elements, e.g. lines, images or blocks.'''
        super().__init__(instances, parent)
        self._floating_image_blocks = []


    def _update_bbox(self, block:Block):
        '''Override. The parent is ``Layout``, which is not necessary to update its bbox. 
        So, do nothing but required here.
        '''
        pass


    @property
    def floating_image_blocks(self):
        return self._floating_image_blocks


    @property
    def lattice_table_blocks(self):
        '''Get lattice table blocks contained in this Collection.'''
        return list(filter(
            lambda block: isinstance(block, Block) and block.is_lattice_table_block, self._instances))


    @property
    def stream_table_blocks(self):
        '''Get stream table blocks contained in this Collection.'''
        return list(filter(
            lambda block: isinstance(block, Block) and block.is_stream_table_block, self._instances))


    @property
    def table_blocks(self):
        '''Get table blocks contained in this Collection.'''
        return list(filter(
            lambda block: isinstance(block, Block) and block.is_table_block, self._instances))


    @property
    def inline_image_blocks(self):
        '''Get inline image blocks contained in this Collection.'''
        return list(filter(
            lambda block: isinstance(block, Block) and block.is_inline_image_block, self._instances))


    @property
    def text_blocks(self):
        '''Get text/image blocks contained in this Collection.'''
        return list(filter(
            lambda block: isinstance(block, Block) and block.is_text_image_block, self._instances))


    def restore(self, raws:list):
        '''Clean current instances and restore them from source dict.
        ImageBlock is converted to ImageSpan contained in TextBlock.

        Args:
            raws (list): A list of raw dicts representing text/image/table blocks.

        Returns:
            Blocks: self
        '''
        self.reset()  # clean current instances
        for raw_block in raws:
            block_type = raw_block.get('type', -1) # type: int
            
            # inline image block -> text block
            if block_type == BlockType.IMAGE.value:
                block = ImageBlock(raw_block).to_text_block()
            
            # text block
            elif block_type == BlockType.TEXT.value:
                block = TextBlock(raw_block)
            
            # table block
            elif block_type in (BlockType.LATTICE_TABLE.value, BlockType.STREAM_TABLE.value):
                block = TableBlock(raw_block)
            
            else:
                block = None            
            
            # add to list
            self.append(block)
        
        return self


    def clean_up(self, float_image_ignorable_gap:float, line_overlap_threshold:float):
        '''Clean up blocks in page level.

        * convert to lines
        * remove lines out of page
        * remove transformed text: text direction is not (1, 0) or (0, -1)
        * remove empty lines

        Args:
            float_image_ignorable_gap (float): Regarded as float image if the intersection exceeds this value.
            line_overlap_threshold (float): remove line if the intersection exceeds this value.

        .. note::
            The block structure extracted from ``PyMuPDF`` might be unreasonable, e.g. 
            * one real paragraph is split into multiple blocks; or
            * one block consists of multiple real paragraphs
        '''
        if not self._instances: return

        # convert to lines
        instances = []
        for block in self._instances:
            if not isinstance(block, (ImageBlock, TextBlock)): continue
            instances.extend(block.lines)
        
        # delete invalid lines
        page_bbox = self.parent.working_bbox
        f = lambda line: line.bbox.intersects(page_bbox) and \
                        not line.white_space_only and (
                            line.is_horizontal_text or line.is_vertical_text)
        instances = list(filter(f, instances))

        # delete redundant blanks
        for line in instances: line.strip()

        # detect floating images
        self.reset(instances) \
            ._identify_floating_images(float_image_ignorable_gap) \
            ._remove_overlapped_lines(line_overlap_threshold)


    def assign_to_tables(self, tables:list):
        '''Add blocks (line or sub-table) to associated cells of given tables.

        Args:
            tables (list): A list of TableBlock instances.
        '''        
        if not tables: return

        # assign blocks to table region        
        blocks_in_tables = [[] for _ in tables] # type: list[list[Line|TableBlock]]
        blocks = []   # type: list[Line|TableBlock]
        for block in self._instances:
            self._assign_block_to_tables(block, tables, blocks_in_tables, blocks)

        # assign blocks to associated cells
        for table, blocks_in_table in zip(tables, blocks_in_tables):
            # no contents for this table
            if not blocks_in_table: continue
            table.assign_blocks(blocks_in_table)

        # sort in natural reading order and update layout blocks
        blocks.extend(tables)
        self.reset(blocks)


    def collect_stream_lines(self, potential_shadings:list, line_separate_threshold:float):
        '''Collect elements in Line level (line or table bbox), which may contained in a stream table region.
        
        Table may exist on the following conditions:

        * blocks in a row don't follow flow layout; or
        * block is contained in potential shading
        
        Args:
            potential_shadings (list): a group of shapes representing potential cell shading
            line_separate_threshold (float): two separate lines if the x-distance exceeds this value
        
        Returns:
            list: A list of Lines. Each group of Lines represents a potential table.
        
        .. note::
            ``PyMuPDF`` may group multi-lines in a row as a text block while each line belongs to different
            cell. So, it's required to deep into line level.
        '''
        # group lines by row: any intersection would be counted as same group
        # NOTE: consider real text direction when all lines have same orientation, i.e. either horizontal or 
        # vertical; otherwise, consider the default direction, i.e. horizontal.
        rows = self.group_by_rows(text_direction=not self.is_mix_text)

        # get sub-lines from block: line or table block
        def sub_line(block):
            return block if isinstance(block, Line) else Line().update_bbox(block.outer_bbox)
        
        # exclude potential shading in white bg-color
        shadings_exclude_white = list(filter(
            lambda shape: not shape.is_determined and shape.color!=rgb_value([1,1,1]), 
            potential_shadings))
        def contained_in_shadings(block):
            for shading in shadings_exclude_white:
                if shading.contains(block, threshold=constants.FACTOR_MOST): return True
            return False

        # store text lines in a potential table
        res = []           # type: list[Lines]
        table_lines = []   # type: list[Line]
        def close_table():
            if not table_lines: return
            res.append(Lines(table_lines))
            table_lines.clear()
        
        
        # check row by row 
        ref_pos = rows[0].bbox.y1
        cell_layout = isinstance(self.parent, Cell)
        for row in rows:

            bbox = row.bbox

            # flow layout or not?
            if not row.is_flow_layout(line_separate_threshold, cell_layout=cell_layout): 
                table_lines.extend([sub_line(block) for block in row])

            else:
                close_table()

            # contained in shading or not?
            for block in row:
                if contained_in_shadings(block): table_lines.append(sub_line(block))
            
            # close table if significant vertical distance 
            if bbox.y0-ref_pos>=50: close_table()
            
            # update reference pos
            ref_pos = bbox.y1

        # don't forget last table
        close_table()

        return res


    def parse_block(self, max_line_spacing_ratio:float, line_break_free_space_ratio:float, new_paragraph_free_space_ratio:float):
        '''Group lines into text block.'''
        # sort in normal reading order
        self.sort_in_reading_order_plus()

        # join lines with similar properties, e.g. spacing, together into text block
        blocks = self._join_lines_vertically(max_line_spacing_ratio)

        # split text block by checking text
        blocks = self._split_text_block_vertically(blocks,
            line_break_free_space_ratio, 
            new_paragraph_free_space_ratio)
        
        self.reset(blocks)


    def parse_text_format(self, rects, delete_end_line_hyphen:bool):
        '''Parse text format with style represented by stroke/fill shapes.
        
        Args:
            rects (Shapes): Potential styles applied on blocks.
            delete_end_line_hyphen (bool): delete hyphen at the end of a line if True.
        '''
        # parse text block style one by one
        for block in filter(lambda e: e.is_text_block, self._instances): 
            block.parse_text_format(rects)

            # adjust word at the end of each line
            block.lines.adjust_last_word(delete_end_line_hyphen)

    
    def parse_spacing(self, *args):
        '''Calculate external and internal space for text blocks:

        - vertical distance between blocks, i.e. paragraph before/after spacing
        - horizontal distance to left/right border, i.e. paragraph left/right indent
        - vertical distance between lines, i.e. paragraph line spacing
        '''
        if not self._instances: return
        self._parse_block_horizontal_spacing(*args)
        self._parse_block_vertical_spacing()
        self._parse_line_spacing()


    def make_docx(self, doc):
        '''Create page based on parsed block structure. 
        
        Args:
            doc (Document, _Cell): The container to make docx content.
        '''
        def make_table(table_block, pre_table):
            # create dummy paragraph if table before space is set
            # - a minimum line height of paragraph is 0.7pt, so ignore before space if less than this value
            # - but tow adjacent tables will be combined automatically, so adding a minimum dummy paragraph is required
            if table_block.before_space>=constants.MIN_LINE_SPACING or pre_table:
                h = lower_round(table_block.before_space, 1) # round(x,1), but to lower bound
                p = doc.add_paragraph()
                reset_paragraph_format(p, line_spacing=Pt(h))

            # new table            
            table = doc.add_table(rows=table_block.num_rows, cols=table_block.num_cols)
            table.autofit = False
            table.allow_autofit  = False
            table_block.make_docx(table)

        pre_table = False
        cell_layout = isinstance(self.parent, Cell)
        for block in self._instances:
            # make paragraphs
            if block.is_text_image_block:                
                # new paragraph
                p = doc.add_paragraph()
                block.make_docx(p)

                pre_table = False # mark block type
            
            # make table
            elif block.is_table_block:
                make_table(block, pre_table)
                pre_table = True # mark block type

                # NOTE: within a cell, there is always an empty paragraph after table,
                # so, delete it right here.
                # https://github.com/dothinking/pdf2docx/issues/76 
                if cell_layout:
                    delete_paragraph(doc.paragraphs[-1])
       
        # NOTE: If a table is at the end of a page, a new paragraph will be automatically 
        # added by the rending engine, e.g. MS Word, which resulting in an unexpected
        # page break. The solution is to never put a table at the end of a page, so add
        # an empty paragraph and reset its format, particularly line spacing, when a table
        # is created.
        for block in self._instances[::-1]:
            # ignore float image block
            if block.is_float_image_block: continue

            # nothing to do if not end with table block
            if not block.is_table_block: break

            # otherwise, add a small paragraph
            p = doc.add_paragraph()
            reset_paragraph_format(p, Pt(constants.MIN_LINE_SPACING)) # a small line height

  
    def plot(self, page):
        '''Plot blocks in PDF page for debug purpose.'''
        for block in self._instances: block.plot(page)                


    # ----------------------------------------------------------------------------------
    # internal methods
    # ----------------------------------------------------------------------------------
    def _identify_floating_images(self, float_image_ignorable_gap:float):
        '''Identify floating image lines and convert to ImageBlock.'''
        # group lines by connectivity
        groups = self.group_by_connectivity(dx=-float_image_ignorable_gap, dy=-float_image_ignorable_gap)
        
        # identify floating images
        for group in filter(lambda group: len(group)>1, groups):
            for line in filter(lambda line: line.image_spans, group):
                float_image = ImageBlock().from_image(line.spans[0])
                float_image.set_float_image_block()
                self._floating_image_blocks.append(float_image)

                # remove the original image line from flow layout by setting empty bbox
                line.update_bbox((0,0,0,0))

        return self

    def _remove_overlapped_lines(self, line_overlap_threshold:float):
        '''Delete overlapped lines. 
        NOTE: Don't run this method until floating images are excluded.
        '''
        # group lines by overlap
        fun = lambda a, b: a.get_main_bbox(b, threshold=line_overlap_threshold)
        groups = self.group(fun)
        
        # delete overlapped lines
        for group in filter(lambda group: len(group)>1, groups):
            # keep only the line with largest area
            sorted_lines = sorted(group, key=lambda line: line.bbox.get_area())
            for line in sorted_lines[:-1]:
                logging.warning('Ignore Line "%s" due to overlap', line.text)
                line.update_bbox((0,0,0,0))

        return self



    @staticmethod
    def _assign_block_to_tables(block, tables:list, blocks_in_tables:list, blocks:list):
        '''Assign block (line or table block) to contained table region ``blocks_in_tables``,
        or out-of-table ``blocks``.'''
        for table, blocks_in_table in zip(tables, blocks_in_tables):
            # fully contained in a certain table with margin
            if table.contains(block, threshold=constants.FACTOR_MAJOR):
                blocks_in_table.append(block)
                break
            
            # not possible in current table, then check next table
            elif not table.bbox.intersects(block.bbox): 
                continue
        
        # Now, this block is out of all table regions
        else:
            blocks.append(block)


    def _join_lines_vertically(self, max_line_spacing_ratio:float):
        '''Create text blocks by merge lines with same properties (spacing, font, size) in 
        vertical direction. At this moment, the block instance is either Line or TableBlock.
        '''
        idx0, idx1 = (1, 3) if self.is_horizontal_text else (0, 2)
        def get_v_bdy(block):
            '''Coordinates of block top and bottom boundaries.'''
            bbox = block.bbox
            return bbox[idx0], bbox[idx1]

        def vertical_distance(block1, block2):
            u0, u1 = get_v_bdy(block1)
            v0, v1 = get_v_bdy(block2)
            return round(v0-u1, 2)
        
        def line_height(line:Line):
            '''The height of line span with most characters.'''
            span = max(line.spans, key=lambda s: len(s.text))
            h = span.bbox.height if line.is_horizontal_text else span.bbox.width
            return round(h, 2)

        def common_vertical_spacing():
            '''Vertical distance with most frequency: a reference of line spacing.'''        
            ref0, ref1 = get_v_bdy(self._instances[0])
            distances = []
            for block in self._instances[1:]:
                y0, y1 = get_v_bdy(block)
                distances.append(round(y0-ref1, 2))
                ref0, ref1 = y0, y1        
            return max(distances, key=distances.count) if distances else 0.0

        # create text block based on lines
        blocks = [] # type: list[TextBlock]
        lines = []  # type: list[Line]
        def close_text_block():
            if not lines: return
            block = TextBlock()
            block.add(lines)
            blocks.append(block)
            lines.clear()

        # check line by line
        ref_dis = common_vertical_spacing()
        for block in self._instances:
            
            # if current is a table block:
            # - finish previous text block; and
            # - add this table block directly 
            if isinstance(block, TableBlock):
                close_text_block()
                blocks.append(block)
            
            # check two adjacent text lines
            else:
                ref_line = lines[-1] if lines else None

                # first line or in same row with previous line: needn't to create new text block
                if not ref_line or ref_line.in_same_row(block):
                    start_new_block = False
                
                # image line: create new text block
                elif block.image_spans or ref_line.image_spans:
                    start_new_block = True
                
                # lower than common line spacing: needn't to create new text block
                elif vertical_distance(ref_line, block)<=ref_dis+1.0 and \
                    ref_dis<=max_line_spacing_ratio*line_height(ref_line):
                    start_new_block = False
                
                else:
                    start_new_block = True
                
                if start_new_block: close_text_block()
                lines.append(block)

        # don't forget last group
        close_text_block()
       
        return blocks


    @staticmethod
    def _split_text_block_vertically(instances:list, line_break_free_space_ratio:float, new_paragraph_free_space_ratio:float):
        '''Split text block into separate paragraph based on punctuation of sentence.

        .. note::
            Considered only normal reading direction, from left to right, from top
            to bottom.
        '''
        blocks = [] # type: list[TextBlock]
        for block in instances:

            # add block if this isn't a text block
            if not block.is_text_block: 
                blocks.append(block)
                continue
            
            # add split blocks if necessary
            lines_list = block.lines.split_vertically_by_text(line_break_free_space_ratio, 
                                                                new_paragraph_free_space_ratio)
            if len(lines_list)==1:
                blocks.append(block)
            else:
                for lines in lines_list:
                    text_block = TextBlock()
                    text_block.add(lines)
                    blocks.append(text_block)

        return blocks


    def _parse_block_horizontal_spacing(self, *args):
        '''Calculate external horizontal space for text blocks, i.e. alignment mode and left spacing 
        for paragraph in docx:
        
            - horizontal block -> take left boundary as reference
            - vertical block   -> take bottom boundary as reference
        '''
        # bbox of blocks
        # - page level, e.g. blocks in top layout
        # - table level, e.g. blocks in table cell
        bbox = self.parent.working_bbox

        for block in self._instances:
            block.parse_horizontal_spacing(bbox, *args)


    def _parse_block_vertical_spacing(self):
        '''Calculate external vertical space for text blocks, i.e. before/after space in docx.
        
        The vertical spacing is determined by the vertical distance to previous block.
        For the first block, the reference position is top margin.

        It's easy to set before-space or after-space for a paragraph with ``python-docx``,
        so, if current block is a paragraph, set before-space for it; if current block is 
        not a paragraph, e.g. a table, set after-space for previous block (generally, 
        previous block should be a paragraph).
        '''
        # bbox of blocks
        # - page level, e.g. blocks in top layout
        # - table level, e.g. blocks in table cell
        bbox = self.parent.working_bbox

        # check text direction for vertical space calculation:
        # - normal reading direction (from left to right)    -> the reference boundary is top border, i.e. bbox[1].
        # - vertical text direction, e.g. from bottom to top -> left border bbox[0] is the reference
        idx = 1 if self.is_horizontal_text else 0

        ref_block = self._instances[0]
        ref_pos = bbox[idx]

        for block in self._instances:

            # NOTE: the table bbox is counted on center-line of outer borders, so a half of top border
            # size should be excluded from the calculated vertical spacing
            if block.is_table_block:
                dw = block[0][0].border_width[0] / 2.0 # use top border of the first cell
            
            else:
                dw = 0.0

            start_pos = block.bbox[idx] - dw
            para_space = start_pos-ref_pos

            # modify vertical space in case the block is out of bottom boundary
            dy = max(block.bbox[idx+2]-bbox[idx+2], 0.0)
            para_space -= dy
            para_space = max(para_space, 0.0) # ignore negative value

            # ref to current (paragraph): set before-space for paragraph
            if block.is_text_block:
                # spacing before this paragraph
                block.before_space = para_space

            # if ref to current (image): set before-space for paragraph
            elif block.is_inline_image_block:
                block.before_space = para_space

            # ref (paragraph/image) to current: set after-space for ref paragraph        
            elif ref_block.is_text_block or ref_block.is_inline_image_block:
                ref_block.after_space = para_space

            # situation with very low probability, e.g. ref (table) to current (table)
            # we can't set before space for table in docx, but the tricky way is to
            # create an empty paragraph and set paragraph line spacing and before space
            else:
                # let para_space>=1 Pt to accommodate the dummy paragraph if not the first block
                block.before_space = max(para_space, int(block!=self._instances[0])*constants.MINOR_DIST)

            # update reference block        
            ref_block = block
            ref_pos = ref_block.bbox[idx+2] + dw # assume same bottom border with top one

        # NOTE: when a table is at the end of a page, a dummy paragraph with a small line spacing 
        # is added after this table, to avoid unexpected page break. Accordingly, this extra spacing 
        # must be remove in other place, especially the page space is run out.
        # Here, reduce the last row of table.
        block = self._instances[-1]
        if block.is_table_block: block[-1].height -= constants.MINOR_DIST


    def _parse_line_spacing(self):
        '''Calculate internal vertical space for text blocks, i.e. paragraph line spacing in docx.

        .. note::
            Run parsing block vertical spacing in advance.
        '''
        def is_exact_line_spacing(block):
            '''check line spacing type based on parsed font metrics of text span:
            exact line spacing if no standard line height is extracted
            '''
            for line in block.lines:
                absent_line_heights = list(not span.is_valid_line_height \
                                        for span in line.spans if isinstance(span, TextSpan))
                if any(absent_line_heights): return True
            return False

        for block in self._instances:
            if not block.is_text_block: continue
            
            if is_exact_line_spacing(block):
                block.line_space_type = 0
                block.parse_exact_line_spacing()
            else:
                block.parse_relative_line_spacing()
