# -*- coding: utf-8 -*-
import os
import json
import logging
from time import perf_counter
from multiprocessing import Pool, cpu_count
import fitz
from docx import Document
from .page.Page import Page
from .page.Pages import Pages

# check PyMuPDF>=1.19.x
if list(map(int, fitz.VersionBind.split("."))) < [1, 19, 0]:
    raise SystemExit("PyMuPDF>=1.19.0 is required for pdf2docx.")

# logging
logging.basicConfig(
    level=logging.INFO, 
    format="[%(levelname)s] %(message)s")


class Converter:
    '''The ``PDF`` to ``docx`` converter.
    
    * Read PDF file with ``PyMuPDF`` to get raw layout data page by page, including text,
      image, drawing and its properties, e.g. boundary box, font, size, image width, height.
    * Analyze layout in document level, e.g. page header, footer and margin.
    * Parse page layout to docx structure, e.g. paragraph and its properties like indentaton, 
      spacing, text alignment; table and its properties like border, shading, merging. 
    * Finally, generate docx with ``python-docx``.
    '''

    def __init__(self, pdf_file:str, password:str=None):
        '''Initialize fitz object with given pdf file path.

        Args:
            pdf_file (str): pdf file path.
            password (str): Password for encrypted pdf. Default to None if not encrypted.
        '''
        # fitz object
        self.filename_pdf = pdf_file
        self.password = str(password or '')
        self._fitz_doc = fitz.Document(pdf_file)

        # initialize empty pages container
        self._pages = Pages()


    @property
    def fitz_doc(self): return self._fitz_doc    

    @property
    def pages(self): return self._pages


    def close(self): self._fitz_doc.close()


    @property
    def default_settings(self):
        '''Default parsing parameters.'''
        return {
            'debug'                          : False,  # plot layout if True
            'ocr'                            : 0,      # ocr status: 0 - no ocr; 1 - to do ocr; 2 - ocr-ed pdf
            'ignore_page_error'              : True,   # not break the conversion process due to failure of a certain page if True
            'multi_processing'               : False,  # convert pages with multi-processing if True
            'cpu_count'                      : 0,      # working cpu count when convert pages with multi-processing
            'min_section_height'             : 20.0,   # The minimum height of a valid section.
            'connected_border_tolerance'     : 0.5,    # two borders are intersected if the gap lower than this value
            'max_border_width'               : 6.0,    # max border width
            'min_border_clearance'           : 2.0,    # the minimum allowable clearance of two borders
            'float_image_ignorable_gap'      : 5.0,    # float image if the intersection exceeds this value
            'page_margin_factor_top'         : 0.5,    # [0,1] reduce top margin by factor
            'page_margin_factor_bottom'      : 0.5,    # [0,1] reduce bottom margin by factor
            'shape_min_dimension'            : 2.0,    # ignore shape if both width and height is lower than this value
            'max_line_spacing_ratio'         : 1.5,    # maximum line spacing ratio: line spacing / line height
            'line_overlap_threshold'         : 0.9,    # [0,1] delete line if the intersection to other lines exceeds this value
            'line_break_width_ratio'         : 0.5,    # break line if the ratio of line width to entire layout bbox is lower than this value
            'line_break_free_space_ratio'    : 0.1,    # break line if the ratio of free space to entire line exceeds this value
            'line_separate_threshold'        : 5.0,    # two separate lines if the x-distance exceeds this value
            'new_paragraph_free_space_ratio' : 0.85,   # new paragraph if the ratio of free space to line height exceeds this value
            'lines_left_aligned_threshold'   : 1.0,    # left aligned if d_x0 of two lines is lower than this value (Pt)
            'lines_right_aligned_threshold'  : 1.0,    # right aligned if d_x1 of two lines is lower than this value (Pt)
            'lines_center_aligned_threshold' : 2.0,    # center aligned if delta center of two lines is lower than this value
            'clip_image_res_ratio'           : 4.0,    # resolution ratio (to 72dpi) when cliping page image
            'min_svg_gap_dx'                 : 15.0,   # merge adjacent vector graphics if the horizontal gap is less than this value
            'min_svg_gap_dy'                 : 2.0,    # merge adjacent vector graphics if the vertical gap is less than this value
            'min_svg_w'                      : 2.0,    # ignore vector graphics if the bbox width is less than this value
            'min_svg_h'                      : 2.0,    # ignore vector graphics if the bbox height is less than this value
            'extract_stream_table'           : False,  # don't consider stream table when extracting tables
            'parse_lattice_table'            : True,   # whether parse lattice table or not; may destroy the layout if set False
            'parse_stream_table'             : True,   # whether parse stream table or not; may destroy the layout if set False
            'delete_end_line_hyphen'         : False   # delete hyphen at the end of a line
        }

    # -----------------------------------------------------------------------
    # Parsing process: load -> analyze document -> parse pages -> make docx
    # -----------------------------------------------------------------------

    def parse(self, start:int=0, end:int=None, pages:list=None, **kwargs):
        '''Parse pages in three steps:
        * open PDF file with ``PyMuPDF``
        * analyze whole document, e.g. page section, header/footer and margin
        * parse specified pages, e.g. paragraph, image and table

        Args:
            start (int, optional): First page to process. Defaults to 0, the first page.
            end (int, optional): Last page to process. Defaults to None, the last page.
            pages (list, optional): Range of page indexes to parse. Defaults to None.
            kwargs (dict, optional): Configuration parameters. 
        '''
        return self.load_pages(start, end, pages) \
            .parse_document(**kwargs) \
            .parse_pages(**kwargs)


    def load_pages(self, start:int=0, end:int=None, pages:list=None):
        '''Step 1 of converting process: open PDF file with ``PyMuPDF``, 
        especially for password encrypted file.
        
        Args:
            start (int, optional): First page to process. Defaults to 0, the first page.
            end (int, optional): Last page to process. Defaults to None, the last page.
            pages (list, optional): Range of page indexes to parse. Defaults to None.
        '''
        logging.info(self._color_output('[1/4] Opening document...'))

        # encrypted pdf ?
        if self._fitz_doc.needs_pass:
            if not self.password:
                raise ConversionException(f'Require password for {self.filename_pdf}.')

            elif not self._fitz_doc.authenticate(self.password):
                raise ConversionException('Incorrect password.')

        # initialize empty pages
        num = len(self._fitz_doc)
        self._pages.reset([Page(id=i, skip_parsing=True) for i in range(num)])

        # set pages to parse
        page_indexes = self._page_indexes(start, end, pages, num)
        for i in page_indexes:
            self._pages[i].skip_parsing = False

        return self
    

    def parse_document(self, **kwargs):
        '''Step 2 of converting process: analyze whole document, e.g. page section,
        header/footer and margin.'''
        logging.info(self._color_output('[2/4] Analyzing document...'))
        
        self._pages.parse(self.fitz_doc, **kwargs)
        return self

    
    def parse_pages(self, **kwargs):
        '''Step 3 of converting process: parse pages, e.g. paragraph, image and table.'''
        logging.info(self._color_output('[3/4] Parsing pages...'))

        pages = [page for page in self._pages if not page.skip_parsing]
        num_pages = len(pages)
        for i, page in enumerate(pages, start=1):
            pid = page.id + 1
            logging.info('(%d/%d) Page %d', i, num_pages, pid)
            try:
                page.parse(**kwargs)
            except Exception as e:
                if not kwargs['debug'] and kwargs['ignore_page_error']:
                    logging.error('Ignore page %d due to parsing page error: %s', pid, e)
                else:
                    raise ConversionException(f'Error when parsing page {pid}: {e}')

        return self


    def make_docx(self, docx_filename=None, **kwargs):
        '''Step 4 of converting process: create docx file with converted pages.
        
        Args:
            docx_filename (str): docx filename to write to.
            kwargs (dict, optional): Configuration parameters. 
        '''
        logging.info(self._color_output('[4/4] Creating pages...'))

        # check parsed pages
        parsed_pages = list(filter(
            lambda page: page.finalized, self._pages
        ))
        if not parsed_pages:
            raise ConversionException('No parsed pages. Please parse page first.')

        # docx file to convert to        
        filename = docx_filename or f'{self.filename_pdf[0:-len(".pdf")]}.docx'
        if os.path.exists(filename): os.remove(filename)

        # create page by page        
        docx_file = Document() 
        num_pages = len(parsed_pages)
        for i, page in enumerate(parsed_pages, start=1):
            if not page.finalized: continue # ignore unparsed pages
            pid = page.id + 1
            logging.info('(%d/%d) Page %d', i, num_pages, pid)
            try:
                page.make_docx(docx_file)
            except Exception as e:
                if not kwargs['debug'] and kwargs['ignore_page_error']:
                    logging.error('Ignore page %d due to making page error: %s', pid, e)
                else:
                    raise MakedocxException(f'Error when make page {pid}: {e}')

        # save docx
        docx_file.save(filename)


    # -----------------------------------------------------------------------
    # Store / restore parsed results
    # -----------------------------------------------------------------------

    def store(self):
        '''Store parsed pages in dict format.'''
        return {
            'filename': os.path.basename(self.filename_pdf),
            'page_cnt': len(self._pages), # count of all pages
            'pages'   : [page.store() for page in self._pages if page.finalized], # parsed pages only
        }


    def restore(self, data:dict):
        '''Restore pages from parsed results.'''
        # init empty pages if necessary
        if not self._pages:
            num = data.get('page_cnt', 100)
            self._pages.reset([Page(id=i, skip_parsing=True) for i in range(num)])
        
        # restore pages
        for raw_page in data.get('pages', []):
            idx = raw_page.get('id', -1)
            self._pages[idx].restore(raw_page)


    def serialize(self, filename:str):
        '''Write parsed pages to specified JSON file.'''
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(json.dumps(self.store(), indent=4))
    

    def deserialize(self, filename:str):
        '''Load parsed pages from specified JSON file.'''
        with open(filename, 'r') as f:
            data = json.load(f)
        self.restore(data)


    # -----------------------------------------------------------------------
    # high level methods, e.g. convert, extract table
    # -----------------------------------------------------------------------

    def debug_page(self, i:int, docx_filename:str=None, debug_pdf:str=None, layout_file:str=None, **kwargs):
        '''Parse, create and plot single page for debug purpose.
        
        Args:
            i (int): Page index to convert.
            docx_filename (str): docx filename to write to.
            debug_pdf (str): New pdf file storing layout information. Default to add prefix ``debug_``.
            layout_file (str): New json file storing parsed layout data. Default to ``layout.json``.
        '''
        # include debug information
        # fitz object in debug mode: plot page layout
        # file path for this debug pdf: demo.pdf -> debug_demo.pdf
        path, filename = os.path.split(self.filename_pdf)
        if not debug_pdf: debug_pdf = os.path.join(path, f'debug_{filename}')
        if not layout_file: layout_file  = os.path.join(path, 'layout.json')
        kwargs.update({
            'debug'         : True,
            'debug_doc'     : fitz.Document(),
            'debug_filename': debug_pdf
        })

        # parse and create docx
        self.convert(docx_filename, pages=[i], **kwargs)
        
        # layout information for debugging
        self.serialize(layout_file)


    def convert(self, docx_filename:str=None, start:int=0, end:int=None, pages:list=None, **kwargs):
        """Convert specified PDF pages to docx file.

        Args:
            docx_filename (str, optional): docx filename to write to. Defaults to None.
            start (int, optional): First page to process. Defaults to 0, the first page.
            end (int, optional): Last page to process. Defaults to None, the last page.
            pages (list, optional): Range of page indexes. Defaults to None.
            kwargs (dict, optional): Configuration parameters. Defaults to None.
        
        Refer to :py:meth:`~pdf2docx.converter.Converter.default_settings` for detail of 
        configuration parameters.
        
        .. note::
            Change extension from ``pdf`` to ``docx`` if ``docx_file`` is None.
        
        .. note::
            * ``start`` and ``end`` is counted from zero if ``--zero_based_index=True`` (by default).
            * Start from the first page if ``start`` is omitted.
            * End with the last page if ``end`` is omitted.
        
        .. note::
            ``pages`` has a higher priority than ``start`` and ``end``. ``start`` and ``end`` works only
            if ``pages`` is omitted.

        .. note::
            Multi-processing works only for continuous pages specified by ``start`` and ``end`` only.
        """
        t0 = perf_counter()
        logging.info('Start to convert %s', self.filename_pdf)
        settings = self.default_settings
        settings.update(kwargs)

        # input check
        if pages and settings['multi_processing']:
            raise ConversionException('Multi-processing works for continuous pages '
                                    'specified by "start" and "end" only.')
        
        # convert page by page
        if settings['multi_processing']:
            self._convert_with_multi_processing(docx_filename, start, end, **settings)
        else:
            self.parse(start, end, pages, **settings).make_docx(docx_filename, **settings)

        logging.info('Terminated in %.2fs.', perf_counter()-t0)        


    def extract_tables(self, start:int=0, end:int=None, pages:list=None, **kwargs):
        '''Extract table contents from specified PDF pages.

        Args:
            start (int, optional): First page to process. Defaults to 0, the first page.
            end (int, optional): Last page to process. Defaults to None, the last page.
            pages (list, optional): Range of page indexes. Defaults to None.
            kwargs (dict, optional): Configuration parameters. Defaults to None.
        
        Returns:
            list: A list of parsed table content.
        '''
        # parsing pages first
        settings = self.default_settings
        settings.update(kwargs)
        self.parse(start, end, pages, **settings)

        # get parsed tables
        tables = []
        for page in self._pages:
            if page.finalized: tables.extend(page.extract_tables(**settings))

        return tables

    
    def _convert_with_multi_processing(self, docx_filename:str, start:int, end:int, **kwargs):
        '''Parse and create pages based on page indexes with multi-processing.

        Reference:

            https://pymupdf.readthedocs.io/en/latest/faq.html#multiprocessing
        '''
        # make vectors of arguments for the processes
        cpu = min(kwargs['cpu_count'], cpu_count()) if kwargs['cpu_count'] else cpu_count()        
        prefix = 'pages' # json file writing parsed pages per process
        vectors = [(i, cpu, start, end, self.filename_pdf, self.password, 
                            kwargs, f'{prefix}-{i}.json') for i in range(cpu)]

        # start parsing processes
        pool = Pool()
        pool.map(self._parse_pages_per_cpu, vectors, 1)
        
        # restore parsed page data
        for i in range(cpu):
            filename = f'{prefix}-{i}.json'
            if not os.path.exists(filename): continue            
            self.deserialize(filename)
            os.remove(filename)
        
        # create docx file
        self.make_docx(docx_filename, **kwargs)


    @staticmethod
    def _parse_pages_per_cpu(vector):
        '''Render a page range of a document.
        
        Args:
            vector (list): A list containing required parameters.
                * 0  : segment number for current process                
                * 1  : count of CPUs
                * 2,3: whole pages range to process
                * 4  : pdf filename
                * 5  : password for encrypted pdf
                * 6  : configuration parameters
                * 7  : json filename storing parsed results
        '''        
        # recreate the arguments
        idx, cpu, s, e, pdf_filename, password, kwargs, json_filename = vector

        # open pdf to get page count: all pages are marked to parse temporarily 
        # since don't know which pages to parse for this moment
        cv = Converter(pdf_filename, password)
        cv.load_pages()

        # the specified pages to process
        e = e or len(cv.fitz_doc)
        all_indexes = range(s, e)
        num_pages = len(all_indexes)

        # page segment processed by this cpu
        m = int(num_pages/cpu)
        n = num_pages % cpu
        seg_size = m + int(idx<n)
        seg_from = (m+1)*idx + min(n-idx, 0)
        seg_to = min(seg_from + seg_size, num_pages)
        page_indexes = [all_indexes[i] for i in range(seg_from, seg_to)]

        # now, mark the right pages
        for page in cv.pages: page.skip_parsing = True
        for i in page_indexes: 
            cv.pages[i].skip_parsing = False

        # parse pages and serialize data for further processing
        cv.parse_document(**kwargs) \
            .parse_pages(**kwargs) \
            .serialize(json_filename)
        cv.close()


    @staticmethod
    def _page_indexes(start, end, pages, pdf_len):
        '''Parsing arguments.'''
        if pages: 
            indexes = [int(x) for x in pages]
        else:
            end = end or pdf_len
            s = slice(int(start), int(end))
            indexes = range(pdf_len)[s]
        
        return indexes

    
    @staticmethod
    def _color_output(msg): return f'\033[1;36m{msg}\033[0m'


class ConversionException(Exception): 
    pass

class MakedocxException(ConversionException): 
    pass