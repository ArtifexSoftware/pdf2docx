# -*- coding: utf-8 -*-

'''docx operation methods based on ``python-docx``.
'''

from docx.shared import Pt
from docx.oxml import OxmlElement, parse_xml, register_element_cls
from docx.oxml.ns import qn, nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from docx.enum.text import WD_COLOR_INDEX
from docx.image.exceptions import UnrecognizedImageError
from docx.table import _Cell
from docx.opc.constants import RELATIONSHIP_TYPE
from .share import rgb_value


# ---------------------------------------------------------
# section and paragraph
# ---------------------------------------------------------
def set_equal_columns(section, num=2, space=0):
    """Set section column count and space. All the columns have same width.

    Args:
        section : ``python-docx`` Section instance.
        num (int): Column count. Defaults to 2.
        space (int, optional): Space between adjacent columns. Unit: Pt. Defaults to 0.
    """
    col = section._sectPr.xpath('./w:cols')[0]
    col.set(qn('w:num'), str(num))
    col.set(qn('w:space'), str(int(20*space))) # basic unit 1/20 Pt


def set_columns(section, width_list:list, space=0):
    """Set section column count and space.

    Args:
        section : ``python-docx`` Section instance.
        width_list (list|tuple): Width of each column.
        space (int, optional): Space between adjacent columns. Unit: Pt. Defaults to 0.
    
    Scheme::

        <w:cols w:num="2" w:space="0" w:equalWidth="0">
            <w:col w:w="2600" w:space="0"/>
            <w:col w:w="7632"/>
        </w:cols>
    """
    cols = section._sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(len(width_list)))    
    cols.set(qn('w:equalWidth'), '0')

    # insert column with width
    # default col exists in cols, so insert new col to the beginning
    for w in width_list[::-1]:
        e = OxmlElement('w:col')
        e.set(qn('w:w'), str(int(20*w)))
        e.set(qn('w:space'), str(int(20*space))) # basic unit 1/20 Pt
        cols.insert(0, e)


def delete_paragraph(paragraph):
    '''Delete a paragraph.

    Reference:    
        https://github.com/python-openxml/python-docx/issues/33#issuecomment-77661907
    '''
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def reset_paragraph_format(p, line_spacing:float=1.05):
    '''Reset paragraph format, especially line spacing.

    Two kinds of line spacing, corresponding to the setting in MS Office Word:

    * line_spacing=1.05: single or multiple
    * line_spacing=Pt(1): exactly
    
    Args:
        p (Paragraph): ``python-docx`` paragraph instance.
        line_spacing (float, optional): Line spacing. Defaults to 1.05.
    
    Returns:
        paragraph_format: Paragraph format.
    '''
    pf = p.paragraph_format
    pf.line_spacing = line_spacing # single by default
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.widow_control = True

    # do not adjust spacing between Chinese and Latin/number
    xml = r'<w:autoSpaceDE {} w:val="0"/>'.format(nsdecls('w'))
    p._p.get_or_add_pPr().insert(0, parse_xml(xml))

    xml = r'<w:autoSpaceDN {} w:val="0"/>'.format(nsdecls('w'))
    p._p.get_or_add_pPr().insert(0, parse_xml(xml))

    return pf


def set_hidden_property(p):
    '''Hide paragraph. This method just sets the paragraph property, while the added text must
    be hided explicitly.

        r = p.add_run()
        r.text = "Hidden"
        r.font.hidden = True

    Args:
        p (Paragraph): python-docx created paragraph.
    '''
    pPr = OxmlElement('w:pPr') # paragraph property
    rPr = OxmlElement('w:rPr') # run property
    v = OxmlElement('w:vanish') # hidden
    rPr.append(v)
    pPr.append(rPr)
    p._p.append(rPr)


# ---------------------------------------------------------
# text properties
# ---------------------------------------------------------
def set_char_scaling(p_run, scale:float=1.0):
    '''Set character spacing: scaling. 
    
    Manual operation in MS Word: Font | Advanced | Character Spacing | Scaling.
    
    Args:
        p_run (docx.text.run.Run): Proxy object wrapping <w:r> element.
        scale (float, optional): scaling factor. Defaults to 1.0.
    '''
    p_run._r.get_or_add_rPr().insert(0, 
        parse_xml(r'<w:w {} w:val="{}"/>'.format(nsdecls('w'), 100*scale)))


def set_char_spacing(p_run, space:float=0.0):
    '''Set character spacing. 
    
    Manual operation in MS Word: Font | Advanced | Character Spacing | Spacing.
    
    Args:
        p_run (docx.text.run.Run): Proxy object wrapping <w:r> element.
        space (float, optional): Spacing value in Pt. Expand if positive else condense. Defaults to 0.0.
    '''
    p_run._r.get_or_add_rPr().insert(0, 
        parse_xml(r'<w:spacing {} w:val="{}"/>'.format(nsdecls('w'), 20*space)))


def set_char_shading(p_run, srgb:int):
    '''Set character shading color, in case the color is out of highlight color scope.
    
    Reference: 
        http://officeopenxml.com/WPtextShading.php
    
    Args:
        p_run (docx.text.run.Run): Proxy object wrapping <w:r> element.
        srgb (int): Color value.
    '''
    # try to set highlight first using python-docx built-in method
    # Here give 6/16 of the valid highlight colors
    color_map = {        
        rgb_value((1,0,0)): WD_COLOR_INDEX.RED,
        rgb_value((0,1,0)): WD_COLOR_INDEX.BRIGHT_GREEN,
        rgb_value((0,0,1)): WD_COLOR_INDEX.BLUE,
        rgb_value((1,1,0)): WD_COLOR_INDEX.YELLOW,
        rgb_value((1,0,1)): WD_COLOR_INDEX.PINK,
        rgb_value((0,1,1)): WD_COLOR_INDEX.TURQUOISE
    }
    if srgb in color_map:
        p_run.font.highlight_color = color_map[srgb]

    # set char shading
    else:
        c = hex(srgb)[2:].zfill(6)
        xml = r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(nsdecls('w'), c)
        p_run._r.get_or_add_rPr().insert(0, parse_xml(xml))


def set_char_underline(p_run, srgb:int):
    '''Set underline and color.
    
    Args:
        p_run (docx.text.run.Run): Proxy object wrapping <w:r> element.
        srgb (int): Color value.
    '''
    c = hex(srgb)[2:].zfill(6)
    xml = r'<w:u {} w:val="single" w:color="{}"/>'.format(nsdecls('w'), c)
    p_run._r.get_or_add_rPr().insert(0, parse_xml(xml))


def add_hyperlink(paragraph, url, text):
    """Create a hyperlink within a paragraph object.

    Reference:

        https://github.com/python-openxml/python-docx/issues/74#issuecomment-215678765

    Args:
        paragraph (Paragraph): ``python-docx`` paragraph adding the hyperlink to.
        url (str): The required url.
        text (str): The text displayed for the url.

    Returns: 
        Run: A Run object containing the hyperlink.
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )
    hyperlink.set(qn('w:history'), '1')

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Create a w:rStyle element, note this currently does not add the hyperlink style as its not in
    # the default template, I have left it here in case someone uses one that has the style in it
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')

    # Join all the xml elements together add add the required text to the w:r element
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    return r


# ---------------------------------------------------------
# image properties
# ---------------------------------------------------------
def add_image(p, image_path_or_stream, width, height):
    ''' Add image to paragraph.
    
    Args:
        p (Paragraph): ``python-docx`` paragraph instance.
        image_path_or_stream (str, bytes): Image path or stream.
        width (float): Image width in Pt.
        height (float): Image height in Pt.
    '''
    docx_span = p.add_run()
    try:
        docx_span.add_picture(image_path_or_stream, width=Pt(width), height=Pt(height))
    except UnrecognizedImageError:
        print('Unrecognized Image.')
        return
    
    # exactly line spacing will destroy image display, so set single line spacing instead
    p.paragraph_format.line_spacing = 1.00


class _CT_Anchor(BaseOxmlElement):
    """
    ``<w:anchor>`` element, container for a floating image.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Return a new ``<wp:anchor>`` element populated with the values passed
        as parameters.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Return a new `wp:anchor` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'                    
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapNone/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % ( nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y) )
        )

register_element_cls('wp:anchor', _CT_Anchor)


def add_float_image(p, image_path_or_stream, width, pos_x=None, pos_y=None):
    '''Add float image behind text.
    
    Args:
        p (Paragraph): ``python-docx`` Paragraph object this picture belongs to.
        image_path_or_stream (str, bytes): Image path or stream.
        width (float): Displaying width of picture, in unit Pt.
        pos_x (float): X-position (English Metric Units) to the top-left point of page valid region
        pos_y (float): Y-position (English Metric Units) to the top-left point of page valid region
    '''
    run = p.add_run()
    # parameters for picture, e.g. id, name
    rId, image = run.part.get_or_add_image(image_path_or_stream)
    cx, cy = image.scaled_dimensions(Pt(width), None)
    shape_id, filename = run.part.next_id, image.filename
    anchor = _CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, Pt(pos_x), Pt(pos_y))
    run._r.add_drawing(anchor)


# ---------------------------------------------------------
# table properties
# ---------------------------------------------------------
def indent_table(table, indent:float):
    '''Indent a table.
    
    Args:
        table (Table): ``python-docx`` Table object.
        indent (float): Indent value, the basic unit is 1/20 pt.
    '''
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(20*indent)) # basic unit 1/20 pt for openxml 
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)


def set_cell_margins(cell:_Cell, **kwargs):
    '''Set cell margins. Provided values are in twentieths of a point (1/1440 of an inch).
    
    Reference: 

        * https://blog.csdn.net/weixin_44312186/article/details/104944773
        * http://officeopenxml.com/WPtableCellMargins.php
    
    Args:
        cell (_Cell): ``python-docx`` Cell instance you want to modify.
        kwargs (dict): Dict with keys: top, bottom, start, end.
        
    Usage::
    
        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)    
    '''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
 
    for m in ['top', 'start', 'bottom', 'end']:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
 
    tcPr.append(tcMar)


def set_cell_shading(cell:_Cell, srgb:int):
    '''Set cell background-color.

    Reference:
        https://stackoverflow.com/questions/26752856/python-docx-set-table-cell-background-and-text-color
    
    Args:
        cell (_Cell): ``python-docx`` Cell instance you want to modify
        srgb (int): RGB color value.
    '''
    c = hex(srgb)[2:].zfill(6)
    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), c)))


def set_cell_border(cell:_Cell, **kwargs):
    '''Set cell`s border.
    
    Reference:
        * https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx
        * https://blog.csdn.net/weixin_44312186/article/details/104944110

    Args:
        cell (_Cell): ``python-docx`` Cell instance you want to modify.
        kwargs (dict): Dict with keys: top, bottom, start, end.

    Usage::
    
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
    '''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existence, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existence, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_vertical_cell_direction(cell:_Cell, direction:str='btLr'):
    '''Set vertical text direction for cell.

    Reference:
        https://stackoverflow.com/questions/47738013/how-to-rotate-text-in-table-cells
    
    Args:
        direction (str): Either "tbRl" (top to bottom) or "btLr" (bottom to top).
    '''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)  # btLr tbRl
    tcPr.append(textDirection)
