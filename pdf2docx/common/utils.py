# -*- coding: utf-8 -*-

import random
from io import BytesIO

import fitz
from fitz.utils import getColorList, getColorInfoList

from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_COLOR_INDEX
from docx.image.exceptions import UnrecognizedImageError


# border margin
DM = 1.0

# inch to point
ITP = 72.0

# tolerant rectangle area
DR = fitz.Rect(-DM, -DM, DM, DM) / 2.0


def is_number(str_number):
    try:
        float(str_number)
    except:
        return False
    else:
        return True


def RGB_component_from_name(name:str=''):
    '''Get a named RGB color (or random color) from fitz predefined colors, e.g. 'red' -> (1.0,0.0,0.0).'''
    # get color index
    if name and name.upper() in getColorList():
        pos = getColorList().index(name.upper())
    else:
        pos = random.randint(0, len(getColorList())-1)
        
    c = getColorInfoList()[pos]
    return (c[1] / 255.0, c[2] / 255.0, c[3] / 255.0)


def RGB_component(srgb:int):
    '''srgb value to R,G,B components, e.g. 16711680 -> (255, 0, 0)'''
    # decimal to hex: 0x...
    s = hex(srgb)[2:].zfill(6)
    return [int(s[i:i+2], 16) for i in [0, 2, 4]]


def RGB_value(rgb:list):
    '''RGB components to decimal value, e.g. (1,0,0) -> 16711680'''
    res = 0
    for (i,x) in enumerate(rgb):
        res += int(x*(16**2-1)) * 16**(4-2*i)
    return int(res)


def CMYK_to_RGB(c:float, m:float, y:float, k:float, cmyk_scale:float=100):
    ''' CMYK components to GRB value.'''
    r = (1.0 - c / float(cmyk_scale)) * (1.0 - k / float(cmyk_scale))
    g = (1.0 - m / float(cmyk_scale)) * (1.0 - k / float(cmyk_scale))
    b = (1.0 - y / float(cmyk_scale)) * (1.0 - k / float(cmyk_scale))
    res = RGB_value((r, g, b)) # type: int
    return res


def to_Highlight_color(sRGB:int):
    '''pre-defined color index for highlighting text with python-docx'''
    # part of pre-defined colors
    color_map = {        
        RGB_value((1,0,0)): WD_COLOR_INDEX.RED,
        RGB_value((0,1,0)): WD_COLOR_INDEX.BRIGHT_GREEN,
        RGB_value((0,0,1)): WD_COLOR_INDEX.BLUE,
        RGB_value((1,1,0)): WD_COLOR_INDEX.YELLOW,
        RGB_value((1,0,1)): WD_COLOR_INDEX.PINK,
        RGB_value((0,1,1)): WD_COLOR_INDEX.TURQUOISE
    }
    return color_map.get(sRGB, WD_COLOR_INDEX.YELLOW)


def get_main_bbox(bbox_1:fitz.Rect, bbox_2:fitz.Rect, threshold:float=0.95):
    ''' If the intersection of bbox_1 and bbox_2 exceeds the threshold, return the union of
        these two bbox-es; else return None.
    '''
    # areas
    b = bbox_1 & bbox_2
    a1, a2, a = bbox_1.getArea(), bbox_2.getArea(), b.getArea()

    # no intersection
    if not b: return fitz.Rect()

    # Note: if bbox_1 and bbox_2 intersects with only an edge, b is not empty but b.getArea()=0
    # so give a small value when they're intersected but the area is zero
    factor = a/min(a1,a2) if a else 1e-6
    if factor >= threshold:
        return bbox_1 | bbox_2
    else:
        return fitz.Rect()


def parse_font_name(font_name):
    '''parse raw font name extracted with pymupdf, e.g.
        BCDGEE+Calibri-Bold, BCDGEE+Calibri
    '''
    font_name = font_name.split('+')[-1]
    font_name = font_name.split('-')[0]
    return font_name


def is_vertical_aligned(bbox1:fitz.Rect, bbox2:fitz.Rect, horizontal:bool=True, factor:float=0.0):
    ''' Check whether two boxes have enough intersection in vertical direction, i.e. perpendicular to reading direction.
        An enough intersection is defined based on the minimum width of two boxes:
        L1+L2-L>factor*min(L1,L2)
        ---
        Args:
        - bbox1, bbox2: bbox region in fitz.Rect type.
        - horizontal  : is reading direction from left to right? True by default.
        - factor      : threshold of overlap ratio, the larger it is, the higher probability the two bbox-es are aligned.

        
    '''
    if not bbox1 or not bbox2:
        return False

    if horizontal: # reading direction: x
        L1 = bbox1.x1-bbox1.x0
        L2 = bbox2.x1-bbox2.x0
        L = max(bbox1.x1, bbox2.x1) - min(bbox1.x0, bbox2.x0)
    else:
        L1 = bbox1.y1-bbox1.y0
        L2 = bbox2.y1-bbox2.y0
        L = max(bbox1.y1, bbox2.y1) - min(bbox1.y0, bbox2.y0)

    res = L1+L2-L>=factor*max(L1,L2) # type: bool

    return res


def is_horizontal_aligned(bbox1:fitz.Rect, bbox2:fitz.Rect, horizontal:bool=True, factor:float=0.0):
    ''' Check whether two boxes have enough intersection in horizontal direction, i.e. the reading direction.
        An enough intersection is defined based on the minimum width of two boxes:
        L1+L2-L>factor*min(L1,L2)
        ---
        Args:
        - bbox1, bbox2: bbox region in fitz.Rect type.
        - horizontal  : is reading direction from left to right? True by default.
        - factor      : threshold of overlap ratio, the larger it is, the higher probability the two bbox-es are aligned.
    '''
    # it is opposite to vertical align situation
    return is_vertical_aligned(bbox1, bbox2, not horizontal, factor)


def in_same_row(bbox1:fitz.Rect, bbox2:fitz.Rect):
    ''' Check whether two boxes are in same row/line:
        - yes: the bottom edge of each box is lower than the centerline of the other one;
        - otherwise, not in same row.

        Note the difference with `is_horizontal_aligned`. They may not in same line, though
        aligned horizontally.
    '''
    if not bbox1 or not bbox2:
        return False

    c1 = (bbox1.y0 + bbox1.y1) / 2.0
    c2 = (bbox2.y0 + bbox2.y1) / 2.0

    # Note y direction under PyMuPDF context
    res = c1<bbox2.y1 and c2<bbox1.y1 # type: bool
    return res


def expand_centerline(start: list, end: list, width:float=2.0):
    ''' convert centerline to rectangle shape.
        centerline is represented with start/end points: (x0, y0), (x1, y1).
    '''
    h = width / 2.0
    x0, y0 = start
    x1, y1 = end

    # consider horizontal or vertical line only
    if x0==x1 or y0==y1:
        res = (x0-h, y0-h, x1+h, y1+h)
    else:
        res = None

    return res


def new_page_section(doc:fitz.Document, width:float, height:float, title):
    '''New page with title shown in page center.'''
    # insert a new page
    page = doc.newPage(width=width, height=height)

    # plot title in page center
    gray = RGB_component_from_name('gray')
    f = 10.0
    page.insertText((width/4.0, (height+height/f)/2.0), title, color=gray, fontsize=height/f)
    return page


def new_page_with_margin(doc:fitz.Document, width:float, height:float, margin:tuple, title:str):
    '''Insert a new page and plot margin borders.'''
    # insert a new page
    page = doc.newPage(width=width, height=height)
    
    # plot borders if page margin is provided
    if margin:
        blue = RGB_component_from_name('blue')
        args = {
            'color': blue,
            'width': 0.5
        }
        dL, dR, dT, dB = margin
        page.drawLine((dL, 0), (dL, height), **args) # left border
        page.drawLine((width-dR, 0), (width-dR, height), **args) # right border
        page.drawLine((0, dT), (width, dT), **args) # top
        page.drawLine((0, height-dB), (width, height-dB), **args) # bottom

    # plot title at the top-left corner
    gray = RGB_component_from_name('gray')
    page.insertText((5, 16), title, color=gray, fontsize=15)
    
    return page


def debug_plot(title:str, plot:bool=True, category:str='layout'):
    ''' Plot layout / shapes for debug mode when the following conditions are all satisfied:
          - plot=True
          - layout has been changed: the return value of `func` is True
          - debug mode: kwargs['debug']=True
          - the pdf file to plot layout exists: kwargs['doc'] is not None        
        ---        
        Args:
          - title: page title
          - plot: plot layout/shape if true
          - category: 
            - 'layout': plot all blocks
            - 'table' : plot explicit table blocks only
            - 'implicit_table' : plot implicit table blocks only
            - 'shape' : plot rectangle shapes
    '''
    def wrapper(func):
        def inner(*args, **kwargs):
            # execute function
            res = func(*args, **kwargs)

            # check if plot layout
            debug = kwargs.get('debug', False)
            doc = kwargs.get('doc', None)
            if plot and res and debug and doc is not None:
                layout = args[0] # assert Layout object
                layout.plot(doc, title, category)
        return inner
    return wrapper


# ------------ docx ---------------

def reset_paragraph_format(p, line_spacing:float=1.05):
    ''' Reset paragraph format, especially line spacing.
        ---
        Args:
          - p: docx paragraph instance
        
        Two kinds of line spacing, corresponding to the setting in MS Office Word:
        - line_spacing=1.05: single or multiple
        - line_spacing=Pt(1): exactly
    '''
    pf = p.paragraph_format
    pf.line_spacing = line_spacing # single by default
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.widow_control = True
    return pf


def add_stop(p, pos:float, current_pos:float):
    ''' Set horizontal position in current position with tab stop.
        ---
        Args: 
          - p: docx paragraph instance
          - pos: target position in Pt
          - current_pos: current position in Pt

        Note: multiple tab stops may exist in paragraph, 
              so tabs are added based on current position and target position.         
    '''
    # ignore small pos
    if pos < Pt(DM): return

    # add tab to reach target position
    for t in p.paragraph_format.tab_stops:
        if t.position < current_pos:
            continue
        elif t.position<pos or abs(t.position-pos)<=Pt(DM):
            p.add_run().add_tab()
        else:
            break


def add_image(p, byte_image, width):
    ''' Add image to paragraph.
        ---
        Args:
          - p: docx paragraph instance
          - byte_image: bytes for image source
          - width: image width
    '''
    # TODO: docx.image.exceptions.UnrecognizedImageError
    docx_span = p.add_run()
    try:
        docx_span.add_picture(BytesIO(byte_image), width=Pt(width))
    except UnrecognizedImageError:
        print('TODO: Unrecognized Image.')


def indent_table(table, indent:float):
    ''' indent table.
        ---
        Args:
          - table: docx table object
          - indent: indent value, the basic unit is 1/20 pt
    '''
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(20*indent)) # basic unit 1/20 pt for openxml 
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)


def set_cell_margins(cell, **kwargs):
    ''' Set cell margins. Provided values are in twentieths of a point (1/1440 of an inch).
        ---
        Args:
          - cell:  actual cell instance you want to modify
          - kwargs: a dict with keys: top, bottom, start, end
        
        Usage:
          - set_cell_margins(cell, top=50, start=50, bottom=50, end=50)
        
        Read more: 
          - https://blog.csdn.net/weixin_44312186/article/details/104944773
          - http://officeopenxml.com/WPtableCellMargins.php
    '''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
 
    for m in ['top', 'start', 'bottom', 'end']:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
 
    tcPr.append(tcMar)


def set_cell_shading(cell, RGB_value):
    ''' set cell background-color.
        ---
        Args:
          - cell:  actual cell instance you want to modify

        https://stackoverflow.com/questions/26752856/python-docx-set-table-cell-background-and-text-color
    '''
    c = hex(RGB_value)[2:].zfill(6)
    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), c)))


def set_cell_border(cell, **kwargs):
    '''
    Set cell`s border.
    
    Reference:
     - https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx
     - https://blog.csdn.net/weixin_44312186/article/details/104944110

    Usage:
    ```
        _set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
    ```
    '''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existence, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existence, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
