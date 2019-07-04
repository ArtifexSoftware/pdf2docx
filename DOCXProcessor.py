'''
create .docx file based on PDF layout data with python package python-docx.
@created: 2019-06-28
@author: train8808@gmail.com
---

typical blocks in layout data and associated docx maker method:
    paragraph block: one or more 'line' sets with region defined with bbox. multi-line sets will be
                     joined in a paragraph which is seperated with tab stop. the position of tab
                     stops could be defined according to bbox.
                     the space after paragraph is defined when processing the next block.
                     an image line is processed by creating a docx paragraph with bbox defining the indent and space.

    table block    : at least two 'line' sets with region defined with bbox. these lines are grouped
                     with left border of bbox and inserted into associated cells of a table. in each
                     cell, the line sets are processed similarly to the lines in paragraph block.

'''


from io import BytesIO
from docx.shared import Pt
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import util



def make_page(doc, layout):
    '''create page based on layout data. To avoid incorrect page break from original document,
       a new page section is created for each page.

       each block has already been marked with 'paragraph' (one line only) or 'table' (multi-
       horizontally aligned lines) type. Avoid overusing table format, seperated table will be
       considered as 'paragraph' before making page.
    '''

    # calculate after-space of paragraph block, set before-space also if previous block is table.
    # besides, mark seperated table blcoks as 'paragraph'
    blocks = layout['blocks']
    num = len(blocks)
    for i,block in enumerate(blocks):
        # find next block in normal reading direction
        for j in range(i+1, num):
            next_block = blocks[j]
            if next_block['lines'][0]['dir']==(1.0, 0.0):
                break
        else:
            next_block = None

        # paragraph1 to paragraph2: set after space for patagraph1
        # table to paragraph: set before space for paragraph
        before_space, after_space = 0.0, 0.0
        if next_block:
            space = max(next_block['bbox'][1]-block['bbox'][3], 0.0)
            if block['type']==0:
                block['after_space'] = space
            elif next_block['type']==0:
                next_block['before_space'] = space

    # new page section
    # a default section is created when initialize the document,
    # so we do not have to add section for the first time.
    if not doc.paragraphs:
        section = doc.sections[0]
    else:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
    width, height = layout['width'], layout['height']
    section.page_width  = Pt(width)
    section.page_height = Pt(height)

    # set page margin
    left,right,top,bottom = layout['margin']    
    section.left_margin = Pt(layout['margin'][0])
    section.right_margin = Pt(layout['margin'][1])
    section.top_margin = Pt(layout['margin'][2])
    section.bottom_margin = Pt(layout['margin'][3])

    # ref_table indicates whether previous block is in table format
    ref_table = None
    for block in blocks:
        # make paragraphs
        if block['type']==0:
            ref_table = None
            if block['lines'][0]['dir'] == (1.0 ,0):
                make_paragraph(doc, block, layout['margin'])
            else:
                make_vertical_paragraph(doc, block)
        # make table
        else:
            ref_table = make_table(doc, ref_table, block, width, layout['margin'])            


def make_paragraph(doc, block, page_margin):
    '''create paragraph for a text block.
       join line sets with TAB and set position according to bbox
    '''
    # new paragraph
    lines = block['lines']
    p = doc.add_paragraph()
    if lines[0]['wmode']==1:
        p.style = 'List Bullet'

    # indent and space setting
    pf = reset_paragraph_format(p)
    pf.space_before = Pt(block.get('before_space', 0.0))
    pf.space_after = Pt(block.get('after_space', 0.0))

    # add text / image: multi-line sets are seperated with TAB
    first = True
    for line in lines:
        pos = line['bbox'][0]-page_margin[0]
        if first:
            pf.left_indent = Pt(pos)
            first = False
        else:
            pf.tab_stops.add_tab_stop(Pt(pos))
            p.add_run('\t')

        # line sets
        add_line(p, line)

    return p
    

def make_vertical_paragraph(doc, block):
    pass


def make_table(doc, table, block, page_width, page_margin):
    '''create table for a text block
       count of columns are checked, combine rows if next block is also in table format
    '''
    lines = block['lines']

    # check count of columns
    boundaries = [round(line['bbox'][0], 2) for line in lines]
    boundaries = list(set(boundaries))
    boundaries.sort()

    # check table
    # create new table if previous object is not a table, or the count of columns is inconsistent
    if not table or len(table.columns)!=len(boundaries):
        table = doc.add_table(rows=1, cols=len(boundaries))
        indent_table(table, 20*(boundaries[0]-page_margin[0])) # basic unit 1/20 pt for openxml 
        cells = table.rows[0].cells

    # otherwise, merge to previous table
    else:
        cells = table.add_row().cells

    # set row height and column width
    table.rows[-1].height = Pt(block['bbox'][3]-block['bbox'][1])

    right_boundaries = boundaries[1:] + [page_width-page_margin[1]]
    for cell, l, r in zip(cells, boundaries, right_boundaries):
        cell.width = Pt(r-l)
    
    # insert into cells
    for cell, x in zip(cells, boundaries):
        cell_lines = list(filter(lambda line: abs(line['bbox'][0]-x)<util.DM, lines))
        first = True
        for line in cell_lines:
            # create paragraph
            # note that a default paragraph is already created
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                s = 'Normal' if line['wmode']!=1 else 'List Bullet'
                p = cell.add_paragraph(style=s)

            # add text/image
            reset_paragraph_format(p)
            add_line(p, line)

    return table

def reset_paragraph_format(p):
    '''paragraph format'''
    pf = p.paragraph_format
    pf.line_spacing = 1 # single
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.widow_control = True
    return pf

def add_line(paragraph, line):
    '''add text/image to a paragraph.       
    '''
    if line['wmode']==2:
        # create image with bytes data stored in block.
        # we could create a temporary image file first and delete it after inserting to word,
        # but a file-like stream is supported and it's more convenient
        span = paragraph.add_run()
        span.add_picture(BytesIO(line['image']), width=Pt(line['bbox'][2]-line['bbox'][0]))
    else:
        text = line['spans']['text']
        flags = line['spans']['flags']
        font = line['spans']['font']
        size = line['spans']['size']

        # line['flags'] is an integer, encoding bools of font properties:
        # bit 0: superscripted (2^0)
        # bit 1: italic (2^1)
        # bit 2: serifed (2^2)
        # bit 3: monospaced (2^3)
        # bit 4: bold (2^4)
        span = paragraph.add_run(text.strip())
        span.italic = bool(flags & 2**1)
        span.bold = bool(flags & 2**4)
        span.font.name = font.split(',')[0] # Calibri,bold => Calibri
        span.font.size = Pt(size)

def indent_table(table, indent):
    '''indent table

       args:
         - indent: indent value, the basic unit is 1/20 pt
    '''
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(indent))
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)
