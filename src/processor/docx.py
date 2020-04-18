'''
create .docx file based on PDF layout data with python package python-docx.
@created: 2019-06-28
@author: train8808@gmail.com
---

'''


from io import BytesIO
from docx.shared import Pt
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from .. import util



def make_page(doc, layout):
    ''' create page based on layout data. 

        To avoid incorrect page break from original document, a new page section
        is created for each page.

        Support general document style only:
          - writing mode: from left to right, top to bottom
          - text direction: horizontal

        The vertical postion of paragraph/table is defined by space_before or 
        space_after property of a paragraph.
    '''
    # new page section
    # a default section is created when initialize the document,
    # so we do not have to add section for the first time.
    if not doc.paragraphs:
        section = doc.sections[0]
    else:
        section = doc.add_section(WD_SECTION.NEW_PAGE)

    width, height = layout['width'], layout['height']
    section.page_width  = Pt(width)
    section.page_height = Pt(height)

    # set page margin
    left,right,top,bottom = layout['margin']    
    section.left_margin = Pt(left)
    section.right_margin = Pt(right)
    section.top_margin = Pt(top)
    section.bottom_margin = Pt(bottom)

    # Calculate vertical space for paragraph block. It'll used when creating paragraph.
    # The vertical position of each block is determined by the vertical distance
    # to previous block. For the first block, the reference position is top margin.
    # It's easy to set before-space or after-space for a paragraph with python-docx,
    # so, if current block is a paragraph, set before-space to previous block;
    # if current block is not a paragraph, e.g. a table, set after-space of previous
    # block (generally, previous block should be a paragraph).
    # 
    ref_block = None 
    for block in layout['blocks']:
        ref_pos = ref_block['bbox'][3] if ref_block else top
        space = max(block['bbox'][1]-ref_pos, 0.0)

        # paragraph-1 (ref) to paragraph-2 (current): set before-space for paragraph-2
        if block['type']==0:
            block['before_space'] = int(space)

        # paragraph (ref) to table (current): set after-space for paragraph
        elif ref_block['type']==0:
            ref_block['after_space'] = int(space)

        # situation with very low probability, e.g. table to table
        else:
            pass

        # update reference block
        ref_block = block

    # ref_table indicates whether previous block is in table format
    ref_table = None
    for block in layout['blocks']:
        # make paragraphs
        if block['type'] in (0, 1):
            ref_table = None
            # horizontal paragraph
            if block['type']==1 or block['lines'][0]['wmode'] == 0:
                make_paragraph(doc, block, width, layout['margin'])
            
            # vertical paragraph
            else:
                make_vertical_paragraph(doc, block)
        
        # make table
        else:
            ref_table = make_table(doc, ref_table, block, width, layout['margin'])            


def make_paragraph(doc, block, width, page_margin):
    '''create paragraph for a text block.
       join line sets with TAB and set position according to bbox.

       Generally, a pdf block is a docx paragraph, with block|line as line in paragraph.
       But without the context, it's not able to recognize a block line as word wrap, or a 
       separate line instead. A rough rule used here:
        - block line will be treated as separate line (append `\n`) by default, except
        - (1) this line and next line are actually in the same line (y-position)
        - (2) if the rest space of this line can't accommodate even one span of next line, 
              it's supposed to be normal word wrap.
    '''
    # new paragraph    
    p = doc.add_paragraph()

    # indent and space setting
    pf = reset_paragraph_format(p)
    pf.space_before = Pt(block.get('before_space', 0.0))
    pf.space_after = Pt(block.get('after_space', 0.0))

    # add image
    if block['type']==1:
        # left indent implemented with tab
        pos = block['bbox'][0]-page_margin[0]
        if abs(pos) > util.DM:
            pf.tab_stops.add_tab_stop(Pt(pos))
            p.add_run().add_tab()
        # create image with bytes data stored in block.
        span = p.add_run()
        span.add_picture(BytesIO(block['image']), width=Pt(block['bbox'][2]-block['bbox'][0]))

    # add text (inline image may exist)
    else:
        for i, line in enumerate(block['lines']):

            # left indent implemented with tab
            pos = line['bbox'][0]-page_margin[0]
            if abs(pos) > util.DM:
                pf.tab_stops.add_tab_stop(Pt(pos))
                p.add_run().add_tab()

            # new line by default
            line_break = True

            # no more lines after last line
            if line==block['lines'][-1]: 
                line_break = False
            
            # different lines in space: y0 of next line should be larger than y1 of current line,
            # otherwise, they are in same line, then don't break the line
            elif block['lines'][i+1]['bbox'][1]<=line['bbox'][3]:
                line_break = False
            
            # now, we have two lines, check whether word wrap or line break
            else:
                # bbox of first span in next line
                x0, _, x1, _ = block['lines'][i+1]['spans'][0]['bbox']
                # word wrap if rest space of this line can't accommodate
                # even one span of next line
                free_space = width-page_margin[1]-line['bbox'][2]
                if x1-x0 >= free_space:
                    line_break = False

            # add line
            for span in line['spans']:
                add_span(span, p, line_break)
    return p
    

def make_vertical_paragraph(doc, block):
    pass


def make_table(doc, table, block, page_width, page_margin):
    '''create table for a text block
       count of columns are checked, combine rows if next block is also in table format
    '''
    lines = block['lines']

    # check count of columns
    boundaries = [round(line['bbox'][0], 2) for line in lines]
    boundaries = list(set(boundaries))
    boundaries.sort()

    # check table
    # create new table if previous object is not a table, or the count of columns is inconsistent
    if not table or len(table.columns)!=len(boundaries):
        table = doc.add_table(rows=1, cols=len(boundaries))
        indent_table(table, 20*(boundaries[0]-page_margin[0])) # basic unit 1/20 pt for openxml 
        cells = table.rows[0].cells

    # otherwise, merge to previous table
    else:
        cells = table.add_row().cells

    # set row height and column width
    table.rows[-1].height = Pt(block['bbox'][3]-block['bbox'][1])

    right_boundaries = boundaries[1:] + [page_width-page_margin[1]]
    for cell, l, r in zip(cells, boundaries, right_boundaries):
        cell.width = Pt(r-l)
    
    # insert into cells
    for cell, x in zip(cells, boundaries):
        cell_lines = list(filter(lambda line: abs(line['bbox'][0]-x)<util.DM, lines))
        first = True
        for line in cell_lines:
            # create paragraph
            # note that a default paragraph is already created
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                s = 'Normal' if line['wmode']!=1 else 'List Bullet'
                p = cell.add_paragraph(style=s)

            # add text/image
            reset_paragraph_format(p)
            add_line(p, line)

    return table


def reset_paragraph_format(p):
    '''paragraph format'''
    pf = p.paragraph_format
    pf.line_spacing = 1 # single
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.widow_control = True
    return pf


def add_span(span, paragraph, break_line=True):
    '''add text span to a paragraph.       
    '''
    # inline image span
    if 'image' in span:
        image_span = paragraph.add_run()
        image_span.add_picture(BytesIO(span['image']), width=Pt(span['bbox'][2]-span['bbox'][0]))

    # text span
    else:
        text_span = paragraph.add_run(span['text'])

        # style setting
        # https://python-docx.readthedocs.io/en/latest/api/text.html#docx.text.run.Font

        # basic font style
        # line['flags'] is an integer, encoding bool of font properties:
        # bit 0: superscripted (2^0)
        # bit 1: italic (2^1)
        # bit 2: serifed (2^2)
        # bit 3: monospaced (2^3)
        # bit 4: bold (2^4)            
        text_span.italic = bool(span['flags'] & 2**1)
        text_span.bold = bool(span['flags'] & 2**4)
        text_span.font.name = util.parse_font_name(span['font'])
        text_span.font.size = Pt(span['size'])
        text_span.font.color.rgb = RGBColor(*util.RGB_component(span['color']))

        # font style parsed from PDF rectangles: 
        # e.g. highlight, underline, strike-through-line
        for style in span.get('style', []):
            t = style['type']
            if t==0:
                text_span.font.highlight_color = util.to_Highlight_color(style['color'])
            elif t==1:
                text_span.font.underline = True
            elif t==2:
                text_span.font.strike = True

    # break line or word wrap?
    if break_line:
        paragraph.add_run('\n')


def indent_table(table, indent):
    '''indent table

       args:
         - indent: indent value, the basic unit is 1/20 pt
    '''
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(indent))
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)
